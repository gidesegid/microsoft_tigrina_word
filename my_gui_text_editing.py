
import ctypes
import sys
import winreg
import os

APP_PROG_ID = "TGFile"
ICON_PATH = os.path.abspath("tgIcon.ico")


# ---------------- ADMIN CHECK ----------------

def is_admin():
    try:
        return ctypes.windll.shell32.IsUserAnAdmin()
    except:
        return False


def relaunch_as_admin():
    ctypes.windll.shell32.ShellExecuteW(
        None,
        "runas",
        sys.executable,
        " ".join([f'"{arg}"' for arg in sys.argv]),
        None,
        1
    )
    sys.exit()


# ---------------- CHECK IF ALREADY REGISTERED ----------------

def tg_already_registered():
    try:
        with winreg.OpenKey(winreg.HKEY_CLASSES_ROOT, ".tg"):
            return True
    except FileNotFoundError:
        return False


# ---------------- REGISTRATION ----------------

def register_tg_file():
    # .tg → TGFile
    with winreg.CreateKey(winreg.HKEY_CLASSES_ROOT, ".tg") as key:
        winreg.SetValue(key, "", winreg.REG_SZ, APP_PROG_ID)

    # TGFile description
    with winreg.CreateKey(winreg.HKEY_CLASSES_ROOT, APP_PROG_ID) as key:
        winreg.SetValue(key, "", winreg.REG_SZ, "TG Document")

    # TGFile icon
    with winreg.CreateKey(
        winreg.HKEY_CLASSES_ROOT,
        rf"{APP_PROG_ID}\DefaultIcon"
    ) as key:
        winreg.SetValue(key, "", winreg.REG_SZ, ICON_PATH)


# ---------------- MAIN BOOTSTRAP ----------------

def ensure_file_type_registered():
    if tg_already_registered():
        return  # ✅ already done, never runs again

    if not is_admin():
        relaunch_as_admin()

    register_tg_file()


import tkinter as tk
from tkinter import ttk, filedialog, colorchooser, messagebox
from tkinter import font as tkFont

from PIL import Image, ImageTk
import win32print
import win32ui
import csv
from docx import Document
import os
from PyPDF2 import PdfReader
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from my_draggable_switch import DraggableSwitch 
from write_in_tigrina import KeyboardListenerHandler
import json
from my_input_text import plain_text_meta_info,extract_ms_word_text_with_metadata
class TextEditor:
    def __init__(self, root,csv_file=None, default_language="english",inputText=None):
        self.root = root
        self.root.title("Tigrina text editor.")
        self.root.geometry("1000x600")
        self.root.iconbitmap("tgIcon.ico")
        
        self.action_info_collection=[]
        self.default_font_family_tools = "Arial"
        self.default_font_size_tools = 12
        self.current_font_family_tools = tk.StringVar(value=self.default_font_family_tools)
        self.current_font_size_tools = tk.IntVar(value=self.default_font_size_tools)
        
        self.keyboard_handler = KeyboardListenerHandler()
        self.keyboard_handler.start_listeners()
        
        self.current_file_in_process=None
        self.current_file = None  # Keep track of current file path
        # Root container for all widgets
        self.main_frame = tk.Frame(self.root)
        self.main_frame.pack(expand=True, fill=tk.BOTH,pady=20)

        self.menu_bar = tk.Menu(self.root)
        self.tools_menu = tk.Menu(self.menu_bar, tearoff=0)

        # Load translations
        
        self.languages, self.translations = self.load_translations(csv_file)
        self.current_language = default_language
        self.toggle_var = tk.BooleanVar(value=True)
        # Language selector at the top-right
        self.language_frame = tk.Frame(self.main_frame, padx=10, pady=1)
        self.language_frame.pack(side=tk.TOP, anchor="ne")  # Top-right corner
        
        self.create_language_selector()

        # Track defined tags
        self.defined_tags = set()
        # Toolbar Frame
        self.toolbar = tk.Frame(self.root, bg="lightgray", padx=5, pady=5)
        self.toolbar.place(x=10)  # Initial position
        self.tool_frame = tk.Frame(self.root, padx=10, pady=5)
        self.tool_frame.pack(side=tk.TOP, fill=tk.X)  # Fixed below the language frame

        # Make the toolbar draggable
        self.toolbar.bind("<Button-1>", self.start_drag)
        self.toolbar.bind("<B1-Motion>", self.do_drag)
        self.font_families = sorted(tkFont.families())  # Get all available font families
        self.font_dropdown = tk.OptionMenu(
            self.toolbar, 
            self.current_font_family_tools, 
            *self.font_families, 
            command=self.update_font
        )
        # Increase dropdown sizes and fonts
        dropdown_font = tkFont.Font(family="Arial", size=12, weight="bold")
        self.font_dropdown.config(font=dropdown_font)
        self.font_dropdown.pack(side=tk.LEFT, padx=5)

        # Font Size Dropdown
        self.font_sizes = [8, 10, 12, 14, 16, 18, 20, 24, 28, 32, 36, 48, 72]  # Predefined sizes
        self.size_dropdown = tk.OptionMenu(
            self.toolbar, 
            self.current_font_size_tools, 
            *self.font_sizes, 
            command=self.update_font
        )
        # Increase dropdown size and font
        dropdown_font = tkFont.Font(family="Arial", size=12, weight="bold")
        self.size_dropdown.config(font=dropdown_font)

        # Adjust the dropdown menu options font
        menu = self.size_dropdown.nametowidget(self.size_dropdown.menuname)
        menu.config(font=dropdown_font)
        self.size_dropdown.pack(side=tk.LEFT, padx=5)

        # Draggable toolbar
        self.create_decorative_toolbar()
        # self.create_function_toolbar()
        # Text area
        self.create_text_area()

        # Create the menu bar
        self.create_menu()

        # Default tag for custom font
        # self.text_area.tag_config("custom_font", font=(self.current_font_family_tools.get(), self.current_font_size_tools.get()))

        
        # Bind typing events to apply the current font
        self.text_area.bind("<KeyPress>", self.apply_new_font)

        if inputText !=None:
          self.openFile(inputText)
       
    def switch_toggled(self, state):
        if state:
            self.keyboard_handler.onActivate()
        else:
            self.keyboard_handler.onDeactivate()
    def load_translations(self, csv_file):
        """Load translations from a CSV file."""
        translations = []
        languages = []
        if csv_file !=None and os.path.exists(csv_file):
            with open(csv_file, mode="r", encoding="utf-8") as file:
                reader = csv.reader(file)
                languages = next(reader)  # First row is the header
                for row in reader:
                    translations.append(row)
        return languages, translations

    def translate(self, text):
        """Translate a text based on the current language."""
        for row in self.translations:
            if row[0] == text:  # Match the English text
                lang_index = self.languages.index(self.current_language)
                return row[lang_index]
        return text  # Fallback to the original text if no translation is found

    def change_language(self, event):
        """Handle language change."""
        language_selector = event.widget
        self.current_language = language_selector.get()
        self.update_ui()
    def update_ui(self):
       
        self.tools_menu.entryconfig(0, label=self.translate("Open")) 
        self.tools_menu.entryconfig(1, label=self.translate("Save")) 
        self.tools_menu.entryconfig(3, label=self.translate("Print Page"))
        self.tools_menu.entryconfig(4, label=self.translate("Insert Image"))
        self.tools_menu.entryconfig(6, label=self.translate("Bold"))
        self.tools_menu.entryconfig(7, label=self.translate("Underline"))
        self.tools_menu.entryconfig(8, label=self.translate("Title"))
        self.tools_menu.entryconfig(9, label=self.translate("Text Color"))
        self.tools_menu.entryconfig(10, label=self.translate("Background Color"))
        self.tools_menu.entryconfig(12, label=self.translate("Indent Right"))
        self.tools_menu.entryconfig(13, label=self.translate("Indent Left"))
        self.tools_menu.entryconfig(15, label=self.translate("3D Text"))
        self.tools_menu.entryconfig(16, label=self.translate("3D Custom"))
        self.tools_menu.entryconfig(18, label=self.translate("Exit"))
    def toggle_listener(self):
        if self.toggle_var.get():  # If the toggle is ON
            self.keyboard_handler.onActivate()
        else:  # If the toggle is OFF
           self.keyboard_handler.onDeactivate()
    def create_language_selector(self):
        """Create the language selector dropdown."""

        # # Create a toggle button
        # toggle_button = tk.Checkbutton(
        #     self.language_frame, 
        #     text="Activate Listener", 
        #     variable=self.toggle_var, 
        #     onvalue=True, 
        #     offvalue=False, 
        #     command=self.toggle_listener,  # Call toggle_listener when toggled
        #     font=("Helvetica", 14)
        # )
        # toggle_button.pack(pady=1)
        # Add the draggable switch
        self.switch = DraggableSwitch(self.language_frame, title="Type Tigrina/ብትግርኛ ምጽሓፍ:", command=self.switch_toggled)
        self.switch.pack()
        language_selector = ttk.Combobox(
            self.language_frame,
            values=self.languages,
            state="readonly"
        )
        language_selector.set(self.current_language)
        language_selector.bind("<<ComboboxSelected>>", self.change_language)
        language_selector.pack() 
    
    def create_text_area(self):
        """Create the text area with scrollbars and A4 paper size behavior."""
        # Define A4 size in pixels (assuming 96 DPI)
        A4_width = 794  # Approximate width in pixels
        A4_height = 1123  # Approximate height in pixels
        
        # Create a frame to hold the text area and scrollbars
        self.text_frame = tk.Frame(self.root, padx=10, pady=10)
        self.text_frame.pack(expand=True, fill=tk.BOTH, side=tk.BOTTOM)
         # Actions Frame for Filter button
        self.actions_frame = tk.Frame(self.root, padx=10, pady=11) 
        self.actions_frame.pack()

        filter_button = tk.Button(self.actions_frame, text="My Filter Here")
        filter_button.pack()

        # New toolbar frame just above the text area
        self.toolbar_frame = tk.Frame(self.text_frame, padx=10, pady=10)
        self.toolbar_frame.pack(fill=tk.X)  # Filling horizontally

        # Add a button to the toolbar
        # toolbar_button = tk.Button(self.toolbar_frame, text="Toolbar Button")
        # toolbar_button.pack(side=tk.LEFT)  # Packing button to the left side of the toolbar frame

        # Add the Collapsing/Expanding Button
        # self.expand_button = tk.Button(self.toolbar_frame, text="Collapse", command=self.toggle_text_area)
        # self.expand_button.pack(side=tk.LEFT)  # Packing button to the left side of the toolbar frame

        # Create the vertical and horizontal scrollbars
        self.vertical_scrollbar = tk.Scrollbar(self.text_frame, orient=tk.VERTICAL)
        self.horizontal_scrollbar = tk.Scrollbar(self.text_frame, orient=tk.HORIZONTAL)

        # # Create the Text widget (the text area)
        self.text_area = tk.Text(
            self.text_frame,
            wrap=tk.WORD,  # Wrap text at word boundaries
            font=(self.default_font_family_tools, self.default_font_size_tools),
            yscrollcommand=self.vertical_scrollbar.set,
            xscrollcommand=self.horizontal_scrollbar.set,
            width=0,  # Allow text area to fill the available space
            height=0,  # Allow height to scale
            bg="white",  # Background color
            padx=30,  # Add padding inside the text area
            pady=20,  # Add padding inside the text area
            undo=True,
        )
         

        # Configure the scrollbars to work with the text area
        self.vertical_scrollbar.config(command=self.text_area.yview)
        self.horizontal_scrollbar.config(command=self.text_area.xview)

        # Pack the scrollbars
        self.vertical_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.horizontal_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)

        # Pack the text area and set its dimensions to A4 size
        self.text_area.pack(expand=True, fill=tk.BOTH)
        self.text_area.config(width=A4_width // 6, height=A4_height // 15)  # Approximate size in terms of number of characters
        
        # Bind text selection event
        # self.text_area.bind("<ButtonRelease-1>", self.apply_font)
        # self.text_area.bind("<Key>", lambda e: self.apply_new_font())
        # Add page breaks and render text
        
        self.render_pages()
    def render_pages(self):
        """Render pages with gray color between them."""
        text_content = self.text_area.get("1.0", "end-1c")  # Get all text content
        lines = text_content.splitlines()
        
        # Define page dimensions
        A4_width = 794  # Width in pixels
        A4_height = 112  # Height in pixels
        
        # Approximate number of lines that fit in one page (based on font size)
        lines_per_page = A4_height // (self.default_font_size_tools + 2)  # Approximation
        chars_per_line = A4_width // (self.default_font_size_tools // 2)  # Approximation for characters per line
        
        # Create pages with text content
        page_lines = []
        current_page = []

        for line in lines:
            while len(line) > chars_per_line:
                # Break lines longer than the page width
                current_page.append(line[:chars_per_line])
                line = line[chars_per_line:]
            current_page.append(line)
            
            # If page exceeds the number of lines allowed per page, create a new page
            if len(current_page) >= lines_per_page:
                page_lines.append(current_page)
                current_page = []
        
        if current_page:  # Add any leftover text to a final page
            page_lines.append(current_page)
        
        # Clear the text area before re-rendering
        self.text_area.delete("1.0", "end")
        
        # Add pages with gray areas between them
        for i, page in enumerate(page_lines):
            if i > 0:  # Add gray space between pages (simulating page break)
                self.text_area.insert("end", "\n" + "-" * 50 + "\n")  # Page separator
            for line in page:
                self.text_area.insert("end", line + "\n")

        # Ensure the text area is scrollable with pages
        self.text_area.yview_moveto(0)  # Ensure we are at the top initially
    def toggle_text_area(self):
        # Get current width of the text area
        current_width = self.text_area.winfo_width()
        
        if current_width > 200:
            # Collapse the text area (set width to a smaller value)
            self.text_area.config(width=20)  # Collapse
            self.expand_button.config(text="Expand")  # Change button text to "Expand"
        else:
            # Expand the text area (set width to a larger value)
            self.text_area.config(width=60)  # Expand
            self.expand_button.config(text="Collapse")  # Change button text to "Collapse"

    def create_decorative_toolbar(self):
        """Create a draggable toolbar with tool symbols/icons."""
        
        # Tool Icons
        tools = [
            ("B", self.make_bold),  # Bold
            ("I", self.make_italic),  # Italic
            ("U", self.make_underline),  # Underline
            ("A", self.change_text_color),  # Text Color
            ("🖍", self.change_background_color),  # Background Color
            ("💧",self.change_selected_text_background_color),
            ("Head 1",lambda:self.make_title(1)),
            ("Head 2",lambda:self.make_title(2)),
            ("•",self.create_bullets),
            ("1. 2. 3.",self.insert_numeric_bullets),
            ("🡐",self.align_left),
            ("🡒",self.align_right),
            ("↔",self.align_center),
            ("↶", self.undo),  # Undo
            ("↷", self.redo),  # Redo
        ]
        
        for symbol, command in tools:
            tool_label = tk.Label(
                self.toolbar,
                text=symbol,
                font=("Helvetica", 16),
                bg="#A8D4AA",
                relief="raised",
                padx=2,
                pady=2,
            )
            tool_label.pack(side=tk.LEFT, padx=3, pady=2)
            tool_label.bind("<Button-1>", lambda event, cmd=command: cmd())  # Bind each tool to its function
    def create_function_toolbar(self):
        """Create a draggable toolbar with tool symbols/icons."""
        # Tool Icons
        # icon = tk.PhotoImage(file="./application_images/filter.png")
        tools = [
            ("Filter", self.meta_info),  # Bold
        ]

        for symbol, command in tools:
            tool_label = tk.Label(
                self.toolbar,
                text=symbol,
                font=("Helvetica", 13),
                bg="#A8D4AA",
                relief="raised",
                padx=2,
                pady=1,
            )
            tool_label.pack(side=tk.LEFT, padx=3, pady=2)
            tool_label.bind("<Button-1>", lambda event, cmd=command: cmd())  # Bind each tool to its function
    
    def meta_info(self):
            if self.current_file_in_process !=None:
                file_extension = os.path.splitext(self.current_file_in_process)
                result=None
                if file_extension.lower() == ".docx":
                    result=extract_ms_word_text_with_metadata(self.current_file_path)
                elif file_extension.lower()==".txt":
                    result=plain_text_meta_info(self.current_file_in_process)
            else:
                pass
    def start_drag(self, event):
        """Record the starting position of the drag."""
        self.drag_start_x = event.x
        self.drag_start_y = event.y
    def do_drag(self, event):
        """Move the toolbar during a drag."""
        x = self.toolbar.winfo_x() - self.drag_start_x + event.x
        y = self.toolbar.winfo_y() - self.drag_start_y + event.y
        self.toolbar.place(x=x, y=y)
    
    
    
    def open_word_document(self):
        # Open a file dialog to select a Word document
        file_path = filedialog.askopenfilename(
            title="Open Word Document",
            filetypes=[("Word Documents", "*.docx")]
        )
        
        if file_path:  # If a file is selected
            try:
                # Read the Word document
                doc = Document(file_path)
                content = ""
                for paragraph in doc.paragraphs:
                    content += paragraph.text + "\n"
                
                # Insert content into the text area
                self.text_area.delete(1.0, tk.END)  # Clear the current text
                self.text_area.insert(tk.END, content)
                self.current_file_path = file_path  # Save the file path for future use
                
                # Optionally update the window title to include the file name
                self.root.title(f"Text Editor - {file_path}")
            
            except Exception as e:
                messagebox.showerror("Error", f"Could not open Word document: {e}")

    def undo(self):
        try:
            self.text_area.edit_undo()
        except tk.TclError:
            pass  # Do nothing if undo is not available
    def redo(self):
        try:
            self.text_area.edit_redo()
        except tk.TclError:
            pass  # Do nothing if redo is not available
   
    def create_menu(self):
        """Create a menu bar with a Tools menu."""
        
        self.tools_menu.add_command(label=self.translate("Open"), command=self.open_file)
        self.tools_menu.add_command(label=self.translate("Save"), command=self.save_file)
        self.tools_menu.add_separator()
        self.tools_menu.add_command(label="Print Page", command=self.hardware_print)
        self.tools_menu.add_command(label="Insert Image", command=self.insert_image)
        self.tools_menu.add_separator()
        self.tools_menu.add_command(label="Bold", command=self.make_bold)
        self.tools_menu.add_command(label="Underline", command=self.make_underline)
        self.tools_menu.add_command(label="Title", command=self.make_title)
        self.tools_menu.add_command(label="Text Color", command=self.change_text_color)
        self.tools_menu.add_command(label="Background Color", command=self.change_background_color)
        self.tools_menu.add_separator()
        self.tools_menu.add_command(label="Indent Right", command=self.indent_right)
        self.tools_menu.add_command(label="Indent Left", command=self.indent_left)
        self.tools_menu.add_separator()
        # self.tools_menu.add_command(label="3D Text", command=self.make_3d)
        # self.tools_menu.add_command(label="3D Custom", command=lambda: self.make_3d_custom(color="blue", offset="2p"))
        self.tools_menu.add_command(label="Save as TG", command=self.save_file_as_tg_word)
        self.tools_menu.add_command(label="Save as word", command=self.save_as_word)
        self.tools_menu.add_command(label="Save as pdf", command=self.save_as_pdf)

        self.tools_menu.add_separator()
        self.tools_menu.add_command(label="Exit", command=self.root.quit)


        self.menu_bar.add_cascade(label="Tools", menu=self.tools_menu)
        self.root.config(menu=self.menu_bar)
    
    def update_font(self, _=None):
        try:
            cursor_postion=self.get_cursor_position()
            self.apply_new_font()
            self.apply_font_to_selection()
        except Exception as error:
            return None
        
    
    def update_font_settings(self, event=None):
        """Update the current font settings based on the dropdowns."""
        self.current_font_family = self.current_font_family_tools.get()
        self.current_font_size = self.current_font_size_tools.get()
    def apply_new_font(self, event=None):
        """Apply the current font settings to the newly entered text."""
        # Get the current cursor position
        current_index = self.text_area.index(tk.INSERT)
        
        # Create a unique tag for the current font settings
        font_tag = f"font_{self.current_font_family_tools.get()}_{str(self.current_font_size_tools.get())}"
        if font_tag not in self.defined_tags:
            self.text_area.tag_config(font_tag, font=(self.current_font_family_tools.get(), self.current_font_size_tools.get()))

        # Apply the unique tag to the newly typed character
        self.text_area.tag_add(font_tag, current_index + " -1c", current_index)
    
    def apply_font_to_selection(self):
        try:
            font_family = self.current_font_family_tools.get()
            font_size = self.current_font_size_tools.get()
            selected_text = self.text_area.tag_ranges(tk.SEL)
            if selected_text:
                self.text_area.tag_add("custom_font", tk.SEL_FIRST, tk.SEL_LAST)
                self.text_area.tag_config("custom_font", font=(font_family, font_size))
            else:
                self.text_area.config(font=(font_family, font_size))
            self.getInfo()
        except tk.TclError:
            return None
    
    def get_all_lines(self):
        # Get the total number of lines in the Text widget
        num_lines = int(self.text_area.index('end-1c').split('.')[0])

        # Retrieve all lines
        lines = []
        for line_num in range(1, num_lines + 1):
            line_text = self.text_area.get(f"{line_num}.0", f"{line_num}.end")
            lines.append(line_text)
        
        return lines
    

    def open_file(self):
         # Open file dialog to select a file
        file_path = filedialog.askopenfilename(
            title="Open File",
            filetypes=[
                ("Tigrina Files", "*.tg"),
                ("Text Files", "*.txt"),
                ("Word Documents", "*.docx"),
                ("PDF Files", "*.pdf"),
                ("All Files", "*.*")
            ]
        )

        if file_path:
            self.openFile(file_path)

    def openFile(self,file_path):
            try:
                # Determine the file extension
                _, file_extension = os.path.splitext(file_path)
                
                if file_extension.lower() == ".docx":
                    self.current_file_in_process=file_path
                    # Handle Word document
                    doc = Document(file_path)
                    content = ""
                    for paragraph in doc.paragraphs:
                        content += paragraph.text + "\n"
                    
                    # Display content in the text area
                    self.text_area.delete(1.0, tk.END)
                    self.text_area.insert(tk.END, content)
                    self.current_file_path = file_path
                    self.current_file = file_path
                
                elif file_extension.lower() == ".txt":
                    self.current_file_in_process=file_path
                    # Handle text file
                    with open(file_path, "r", encoding="utf-8") as file:
                        content = file.read()
                    
                    # Display content in the text area
                    self.text_area.delete(1.0, tk.END)
                    self.text_area.insert(tk.END, content)
                    self.current_file_path = file_path
                    self.current_file = file_path
                elif file_extension.lower() == ".pdf":
                    self.current_file_in_process=file_path
                    # Handle PDF file
                    reader = PdfReader(file_path)
                    content = ""
                    for page in reader.pages:
                        content += page.extract_text() + "\n"
                    
                    # Display content in the text area
                    self.text_area.delete(1.0, tk.END)
                    self.text_area.insert(tk.END, content)
                    self.current_file_path = file_path
                    self.current_file = file_path
                elif file_extension.lower()==".tg":
                    self.open_tg_file(file_path)
                else:
                    # Unsupported file type
                    messagebox.showerror("Unsupported File", "Please select a .txt, .docx, or .pdf file.")
                    return

                # Optionally update the window title to include the file name
                self.root.title(f"Text Editor - {os.path.basename(file_path)}")
            
            except Exception as e:
                messagebox.showerror("Error", f"Could not open the file: {e}")
    
    def save_file(self):
        if self.current_file:
            print("self.current_file vvv:",self.current_file)
            if ".tg" in self.current_file:
                self.save_file_tg()
            else:
                with open(self.current_file, "w") as file:
                    file.write(self.text_area.get("1.0", tk.END))
        else:
            self.save_file_as()
    

    def save_file_as(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".txt",
                                                 filetypes=[("Text Files", "*.txt"), ("All Files", "*.*")])
        if file_path:
            self.current_file = file_path
            self.save_file()  # Save immediately after choosing
    
    def save_file_as_tg_word(self):
        file_path = filedialog.asksaveasfilename(
            defaultextension=".tg",
            filetypes=[
                ("TG Files", "*.tg")
            ]
        )
    
        if file_path:
            self.current_file = file_path
            self.save_file_tg()

    def save_file_tg(self):
        if not self.current_file:
            self.save_file_as()
            return

        # 1. Get text content
        text_content = self.text_area.get("1.0", "end-1c")
        metadata = {
            "action_info_collection": self.action_info_collection
        }
        # 3. Write everything to file
        with open(self.current_file, "w", encoding="utf-8") as file:
            file.write(text_content)
            file.write("\n\n<<<TG_META>>>\n")
            file.write(json.dumps(metadata, indent=4))
    
    def open_tg_file(self, file_path):
        with open(file_path, "r", encoding="utf-8") as file:
            content = file.read()

        if "\n<<<TG_META>>>\n" not in content:
            return

        text_part, meta_part = content.split("\n<<<TG_META>>>\n", 1)

        self.text_area.delete("1.0", tk.END)
        self.text_area.insert("1.0", text_part)

        metadata = json.loads(meta_part)
        action_info_collection = metadata.get("action_info_collection", [])

        for item in action_info_collection:
            start = item["start_index"]
            end = item["end_index"]

            family = item.get("font_family", "Arial")
            size = item.get("font_size", 12)
            tags = item.get("tags", [])
            color = item.get("color")  # ← REAL COLOR VALUE

            weight = "bold" if "bold" in tags else "normal"
            underline = 1 if "underline" in tags else 0

            # ---------------- FONT TAG ----------------
            font_tag = f"font_{family}_{size}_{weight}_{underline}"

            if font_tag not in self.text_area.tag_names():
                text_font = tkFont.Font(
                    family=family,
                    size=size,
                    weight=weight,
                    underline=underline
                )
                self.text_area.tag_configure(font_tag, font=text_font)

            self.text_area.tag_add(font_tag, start, end)

            # ---------------- COLOR TAG ----------------
            if color:
                color_tag = f"fg_{color}"

                if color_tag not in self.text_area.tag_names():
                    self.text_area.tag_configure(color_tag, foreground=color)

                self.text_area.tag_add(color_tag, start, end)

        self.current_file = file_path
        self.action_info_collection = action_info_collection

    def apply_style_to_selection(self, style_name):
        """Apply a style tag (e.g., bold, italic) to the selected text."""
        try:
            # Get the start and end indices of the selected text
            start_index = self.text_area.index(tk.SEL_FIRST)
            end_index = self.text_area.index(tk.SEL_LAST)

            # Add the style tag
            self.text_area.tag_add(style_name, start_index, end_index)
        except tk.TclError:
            # No text selected
            print("No text selected to apply style.")
            return None

    def get_selected_text_with_style(self):
        try:
            start_index = self.text_area.index(tk.SEL_FIRST)
            end_index = self.text_area.index(tk.SEL_LAST)

            selected_text = self.text_area.get(tk.SEL_FIRST, tk.SEL_LAST)

            tags = self.text_area.tag_names(tk.SEL_FIRST)

            color = None
            for tag in tags:
                if tag == 'sel':
                    continue  # Ignore selection tag

                fg = self.text_area.tag_cget(tag, "foreground")
                if fg:
                    color = fg
                    break

            return {
                "text": selected_text,
                "start_index": start_index,
                "end_index": end_index,
                "tags": tags,
                "font_family": self.current_font_family_tools.get(),
                "font_size": self.current_font_size_tools.get(),
                "color": color
            }

        except tk.TclError:
            return None
    
    def add_or_replace_action_info(self, info):
        # The object to be added or replaced
        obj = {
            "text": info["text"],
            "start_index": info["start_index"],
            "end_index": info["end_index"],
            "tags": info['tags'],
            "font_family": info['font_family'],
            "font_size": info['font_size'],
            "color": info['color']
        }
        
        # Search for an object with the same "text" key in the collection
        for i, existing_obj in enumerate(self.action_info_collection):
            if existing_obj["text"] == obj["text"]:
                # Replace the existing object if the "text" matches
                self.action_info_collection[i] = obj
                return
        
        # If no match is found, append the new object
        self.action_info_collection.append(obj)

    def getInfo(self):
        info = self.get_selected_text_with_style()
        print("info>>>>>>:",info)
        if info:
            obj={
                "text":info["text"],
                "start_index":info["start_index"],
                "end_index":info["end_index"],
                "tags":info['tags'],
                "font_family":info['font_family'],
                "font_size":info['font_size'],
                "font_size":info['font_size'],
                "color":info["color"]
            }
            self.add_or_replace_action_info(obj)
        else:
            print("No text selected.")
        cursor_info = self.get_cursor_position()
        print("action_info_collection:",self.action_info_collection)
        
    def get_cursor_position(self):
        cursor_position = self.text_area.index(tk.INSERT)
        line, char = cursor_position.split('.')
        return {
            "line": int(line),      # Line number (1-based)
            "character": int(char)  # Character offset on the line (0-based)
        }

    def hardware_print(self):
        text_to_print = self.text_area.get("1.0", tk.END).strip()
        if not text_to_print:
            return  # Do nothing if there's no content to print

        printer_name = win32print.GetDefaultPrinter()
        hprinter = win32print.OpenPrinter(printer_name)
        try:
            hprinter_job = win32print.StartDocPrinter(hprinter, 1, ("Text Editor Print Job", None, "RAW"))
            win32print.StartPagePrinter(hprinter)
            win32print.WritePrinter(hprinter, text_to_print.encode('utf-8'))
            win32print.EndPagePrinter(hprinter)
            win32print.EndDocPrinter(hprinter)
        except Exception as e:
            print(f"Error printing: {e}")
        finally:
            win32print.ClosePrinter(hprinter)
    def insert_image(self):
        file_path = filedialog.askopenfilename(filetypes=[("Image Files", "*.png;*.jpg;*.jpeg;*.bmp;*.gif"), ("All Files", "*.*")])
        if file_path:
            img = Image.open(file_path)
            img.thumbnail((200, 200))
            photo = ImageTk.PhotoImage(img)

            def on_drag(event):
                image_label.place(x=event.x_root, y=event.y_root)

            image_label = tk.Label(self.text_area, image=photo)
            image_label.image = photo
            image_label.place(relx=0.5, rely=0.5, anchor=tk.CENTER)
            image_label.bind("<B1-Motion>", on_drag)

    def save_as_pdf(self):
        # Get the content of the text area
        text_content = self.text_area.get("1.0", tk.END)
        
        # Open a file dialog to choose where to save the PDF
        file_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF files", "*.pdf")])
        
        if file_path:
            # Create a PDF object using reportlab
            c = canvas.Canvas(file_path, pagesize=letter)
            width, height = letter  # Default letter size
            
            # Define the starting position for the text
            x = 72  # 1 inch from the left
            y = height - 72  # 1 inch from the top
            
            # Set the font for the PDF
            c.setFont("Helvetica", 10)
            
            # Split the content into lines and add it to the PDF
            for line in text_content.splitlines():
                c.drawString(x, y, line)
                y -= 12  # Move down by 12 units for the next line
                
                if y < 72:  # If the content goes below the bottom margin, create a new page
                    c.showPage()
                    c.setFont("Helvetica", 10)
                    y = height - 72
            
            # Save the PDF
            c.save()
    def save_as_word(self):
        # Get the content of the text area
        text_content = self.text_area.get("1.0", tk.END)
        
        # Open a file dialog to choose where to save the Word document
        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        
        if file_path:
            # Create a new Document
            doc = Document()
            
            # Add the text content to the document
            for line in text_content.splitlines():
                doc.add_paragraph(line)
            
            # Save the Word document
            doc.save(file_path)
    # Text Styling
    def make_bold(self):
        current_tags = self.text_area.tag_names("sel.first")
        print("current_tags:",current_tags)
        if "bold" in current_tags:
            self.text_area.tag_remove("bold", "sel.first", "sel.last")
        else:
            self.text_area.tag_add("bold", "sel.first", "sel.last")
            self.text_area.tag_config("bold", font=(self.default_font_family_tools, self.default_font_size_tools, "bold"))
        self.getInfo()
    def make_italic(self):
        current_tags = self.text_area.tag_names("sel.first")
        if "italic" in current_tags:
            self.text_area.tag_remove("italic", "sel.first", "sel.last")
        else:
            self.text_area.tag_add("italic", "sel.first", "sel.last")
            self.text_area.tag_config("italic", font=(self.current_font_family_tools.get(), self.current_font_size_tools.get(), "italic"))
        # info = self.get_selected_text_info()
        self.getInfo()
    def make_underline(self):
        current_tags = self.text_area.tag_names("sel.first")
        if "underline" in current_tags:
            self.text_area.tag_remove("underline", "sel.first", "sel.last")
        else:
            self.text_area.tag_add("underline", "sel.first", "sel.last")
            self.text_area.tag_config("underline", font=(self.current_font_family_tools.get(), self.current_font_size_tools.get(), "underline"))
        self.getInfo()
    def make_title(self,level):
        current_tags = self.text_area.tag_names("sel.first")  # No more AttributeError!
        tag_name = f"title{level}"
        if tag_name in current_tags:
            self.text_area.tag_remove(tag_name, "sel.first", "sel.last")
        else:
            print("self.current_font_size_tools:",self.current_font_size_tools.get())
            base_size = self.current_font_size_tools.get()  # Example base size
            title_size = base_size + (4 - level)
            title_size = max(8, title_size)

            self.text_area.tag_add(tag_name, "sel.first", "sel.last")
            self.text_area.tag_config(tag_name, font=("Arial", title_size, "bold"))
        self.getInfo()

    def make_title2(self):
        current_tags = self.text_area.tag_names("sel.first")
        if "title2" in current_tags:
            self.text_area.tag_remove("title2", "sel.first", "sel.last")
        else:
            self.text_area.tag_add("title2", "sel.first", "sel.last")
            self.text_area.tag_config("title2", font=(self.default_font_family_tools, self.default_font_size_tools + 2, "bold"))
        self.getInfo()
    def change_text_color(self):
        color_code = colorchooser.askcolor(title="Choose text color")[1]
        if color_code:
            self.text_area.tag_add("color", "sel.first", "sel.last")
            self.text_area.tag_config("color", foreground=color_code)
        self.getInfo()
    def change_selected_text_background_color(self, color="yellow"):
        """Change the background color of the selected text."""
        try:
            # Check if there's a selection
            selected_text_start = self.text_area.index("sel.first")
            selected_text_end = self.text_area.index("sel.last")
            color_code = colorchooser.askcolor(title="Choose background color")[1]
            if color_code:
                # self.text_area.configure(bg=color_code)
                 # Define or update the tag for background color
                self.text_area.tag_add("highlight", "sel.first", "sel.last")
                self.text_area.tag_config("highlight", background=color_code)
            print(f"Background color set to {color} for selection: {selected_text_start} to {selected_text_end}")
            self.getInfo()  # Update text info if required
        except tk.TclError:
            # Handle case where no text is selected
            print("No text selected to apply background color.")
            messagebox.showinfo("Select text","No text selected to apply background color.")
    
    def change_background_color(self):
        color_code = colorchooser.askcolor(title="Choose background color")[1]
        if color_code:
            self.text_area.configure(bg=color_code)
        self.getInfo()
   
    def indent_right(self):
        self.text_area.tag_add("indent-right", "sel.first", "sel.last")
        self.text_area.tag_config("indent-right", lmargin1=20, lmargin2=20)
        self.getInfo()
    def indent_left(self):
        self.text_area.tag_add("indent-left", "sel.first", "sel.last")
        self.text_area.tag_config("indent-left", lmargin1=0, lmargin2=0)
        self.getInfo()

    def insert_numeric_bullets(self):
        # Get the current cursor position
        cursor_position = self.text_area.index(tk.INSERT)
        
        # Get the line number at the cursor position
        line_number = int(cursor_position.split('.')[0])
        
        # Get the current selection (if any)
        selected_text = self.text_area.get("sel.first", "sel.last") if self.text_area.tag_ranges("sel") else None
        
        # If there's no selection, we add the bullet to the current line
        if not selected_text:
            selected_text = self.text_area.get(f"{line_number}.0", f"{line_number}.end")
        
        # Split the selected text by lines
        lines = selected_text.split("\n")
        
        # Initialize the line number to start the bullets
        bullet_number = 1
        
        for line in lines:
            # Remove extra spaces at the start of the line
            clean_line = line.strip()
            
            if clean_line:  # Only add bullets to non-empty lines
                # Insert the numeric bullet and the text
                self.text_area.insert(f"{line_number}.0", f"{bullet_number}. {clean_line}\n")
                bullet_number += 1  # Increase the bullet number for the next line
            
            # Move the line number to the next line
            line_number += 1

        # Remove the original selection if it exists
        if selected_text:
            self.text_area.delete("sel.first", "sel.last")
    def create_bullets(self):
        try:
            start_index = self.text_area.index("sel.first")
            end_index = self.text_area.index("sel.last")
        except tk.TclError:
            messagebox.showinfo("Selection Required", "Please select text to create a bullet list.")
            return

        # Add a bullet (*) to the beginning of each selected line
        lines = self.text_area.get(start_index, end_index).split("\n")
        bullet_lines = ["* " + line if not line.startswith("* ") else line for line in lines]
        self.text_area.delete(start_index, end_index)
        self.text_area.insert(start_index, "\n".join(bullet_lines))
    def align_left(self):
        try:
            start_index = self.text_area.index("sel.first")
            end_index = self.text_area.index("sel.last")
        except tk.TclError:
            start_index = "1.0"
            end_index = tk.END

        # Align text to the left
        lines = self.text_area.get(start_index, end_index).split("\n")
        aligned_lines = [line.lstrip() for line in lines]
        self.text_area.delete(start_index, end_index)
        self.text_area.insert(start_index, "\n".join(aligned_lines))
    def align_right(self):
        try:
            start_index = self.text_area.index("sel.first")
            end_index = self.text_area.index("sel.last")
        except tk.TclError:
            start_index = "1.0"
            end_index = tk.END

        # Align text to the right
        lines = self.text_area.get(start_index, end_index).split("\n")
        max_width = max(len(line) for line in lines)
        aligned_lines = [line.rjust(max_width) for line in lines]
        self.text_area.delete(start_index, end_index)
        self.text_area.insert(start_index, "\n".join(aligned_lines))
    def align_center(self):
        try:
            start_index = self.text_area.index("sel.first")
            end_index = self.text_area.index("sel.last")
        except tk.TclError:
            start_index = "1.0"
            end_index = tk.END

        # Align text to the center
        lines = self.text_area.get(start_index, end_index).split("\n")
        max_width = max(len(line) for line in lines)
        aligned_lines = [line.center(max_width) for line in lines]
        self.text_area.delete(start_index, end_index)
        self.text_area.insert(start_index, "\n".join(aligned_lines))


# Initialize the application
if __name__ == "__main__":
    ensure_file_type_registered()
    csv_file = "language.csv"
    root = tk.Tk()
    app = TextEditor(root,csv_file, default_language="english")
    root.mainloop()

