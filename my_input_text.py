import re
import unicodedata
from docx import Document
import os
import spacy
from transformers import T5Tokenizer, T5ForConditionalGeneration,pipeline,AutoTokenizer,AutoModelForCausalLM
import random
import string
from bs4 import BeautifulSoup
import requests
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from PIL import Image, ImageDraw, ImageFont
from PyPDF2 import PdfReader
from common import check_file_type,combine_path,number_to_words,currency_with_number_to_words,is_currency_number,is_number,separate_text_number
import json
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.text_rank import TextRankSummarizer

# import nltk
# nltk.download('wordnet')
import nltk
# nltk.download('punkt')
# nltk.download('punkt_tab')
from nltk.corpus import wordnet
from nltk.stem import WordNetLemmatizer

from deep_translator import GoogleTranslator
from keybert import KeyBERT

import warnings
warnings.filterwarnings("ignore", category=UserWarning)
# Step 3: Extract keywords
kw_model = KeyBERT()

# Load the SpaCy English model
# nlp = spacy.load("en_core_web_sm")
# Initialize WordNet Lemmatizer
lemmatizer = WordNetLemmatizer()

summarizer_bart_facebook = pipeline("summarization", model="facebook/bart-large-cnn")
lineExcludes={'_', '-', ' ', '\n', '\t',''}
englishAbbrivationInTigrina={"a":"ኤ","b":"ቢ","c":"ሲ","d":"ዲ","e":"ኢ","f":"ኤፍ","g":"ጂ","h":"ኤች",
     "i":"ኣይ","j":"ጀይ","k":"ከይ","l":"ኤል","m":"ኤም","n":"ኤን","o":"ኦ","p":"ፒ",
    "q":"ኪው","r":"ኣር","s":"ኤስ","t":"ቲ","u":"ዩ","v":"ቪ","w":"ዳብሊው","x":"ኤክስ",
    "y":"ዋይ","z":"ዜድ",
    "A":"ኤ","B":"ቢ","C":"ሲ","D":"ዲ","E":"ኢ","F":"ኤፍ","G":"ጂ","H":"ኤች",
     "I":"ኣይ","J":"ጀይ","K":"ከይ","L":"ኤል","M":"ኤም","N":"ኤን","O":"ኦ","P":"ፒ",
    "Q":"ኪው","R":"ኣር","S":"ኤስ","T":"ቲ","U":"ዩ","V":"ቪ","W":"ዳብሊው","X":"ኤክስ",
    "Y":"ዋይ","Z":"ዜድ"}

english_tigrina_exceptions={"Google":"ጎግል","Google sheet":"ጎግል ሽት","linkedIn":"ሊንክድኢን"}
customWordToTigrigna = {
    "USA": "ኤሜሪካ",
    "UN": "ተባባሪ መንግስታት",
    "WHO": "ዓለም ኣገልግሎት ሕክምና",
    "UK": "እንግሊዝ",
    "Google":"ጎግል",
    "Google Sheet":"ጎግል ሽት", #Google Sheet
    "Upwork":"ኣፕ ዎርክ", 
    "TopTel":"ቶፕተል",
    "LinkedIn":"ሊንክድ ኢን",
    "Tags":"ታግስ",
    "keywords":"ኣምራት",
    "MDN Web Docs (by Mozilla)":"ኤምድን ወብ ዶክስ",
    "W3Schools":"ዳውብሊው3ስኩል",
    "FreeCodeCamp":"ፍሪኮድካምፕ",
    "Codecademy HTML Course":"ኮድካደምይ",
    "React":"ሪኣክት",
    "Angular":"ኣንጉላር",
    "Vue":"ቪው"
    # Add more words as needed
}
tigrinaNegatives=['ኣይ',    'ዘይ',    'ከይ','ኸይ',    'ንዘይ',   'ንከይ',   'ንኸይ','ኣብዘይ','ብዘይ', 'ብዘይም',  'ዘይም','ምስዘይ','ከምዘይ','እንተዘይ', 'ኣምበይ','ኣምበይም']
def isTextNegativeOrPostive(text):
    try:
        sentiment_pipeline = pipeline("sentiment-analysis")
        result = sentiment_pipeline(text)
        return result[0]["label"]
    except Exception as error:
        return f"Error: There was error investigating if sentence is positive or negative.{error}"
def isTigrinaTextNegativeOrPositive(text):
    return "negative" if any(neg in text for neg in tigrinaNegatives) else "positive"

def convert_abbreviation_to_tigrigna(word):
    match = re.match(r'([A-Za-z]{2,})([.,!?]?)', word)
    if match:
        original = match.group(1)
        punct = match.group(2)

        # Try custom dictionary (case-sensitive or normalize as needed)
        if original in customWordToTigrigna:
            return customWordToTigrigna[original] + punct
        elif original.upper() in customWordToTigrigna:
            return customWordToTigrigna[original.upper()] + punct

        # Fallback: letter-by-letter conversion
        letters = original.upper()
        if all(char in englishAbbrivationInTigrina for char in letters):
            return ''.join(englishAbbrivationInTigrina[char] for char in letters) + punct

    return word

def replace_english_words_with_tigrigna(text):
    # Only proceed if there's English in the text
    if not re.search(r'\b[A-Za-z]{2,}\b', text):
        return text

    def replacer(match):
        word = match.group()
        return convert_abbreviation_to_tigrigna(word)

    return re.sub(r'\b[A-Za-z]{2,}[.,!?]?\b', replacer, text)

def extract_titles_spacy(filename):
    # Read the entire text from the file
    with open(filename, 'r', encoding='utf-8') as file:
        text = file.read()
    doc = nlp(text)
    titles = []
    
    # Iterate through each sentence in the document
    for sent in doc.sents:
        # Check if the sentence is short and has title-like features
        if is_title_by_spacy(sent.text):
            titles.append(sent.text.strip())
    
    return titles

def is_title_by_spacy(text):
    # Simple heuristics to determine if a sentence is a title
    words = text.split()
    if len(words) <= 10 and any(char.isupper() for char in text) and text[-1] not in ".!?":
        return True
    return False

filler_words = [
    "hmm", "uh", "um", "well", "like", "you know", "I mean", "hold on", 
    "oh", "yeah", "okay", "right", "ow", "alright", "yep", "actually", 
    "basically", "literally", "sort of", "kind of", "just", "anyway", 
    "so", "now", "ah", "uh-huh", "gotcha", "see", "that said", "to be honest", 
    "if you will", "as I was saying", "let's see", "if I'm being honest", 
    "I guess", "maybe", "probably", "certainly", "I suppose", "you know what I mean", 
    "for sure", "what I'm saying is", "in other words", "like I said", "to tell the truth", 
    "at the end of the day", "well, you see", "you know what", "that being said", 
    "it's like", "well, actually", "I mean to say", "here's the thing", "just a second", 
    "give me a second", "hold up", "just saying", "let me think", "let me see", 
    "you understand", "what I mean is", "let's be real", "I mean really", "that's the thing", 
    "if I'm honest", "kinda", "sorta", "it's like this", "the thing is", "you know what I mean?", 
    "look", "right now", "like I mentioned", "on the other hand", "to put it another way", 
    "basically", "you feel me", "what's it called", "oh my", "uh-oh", "you know what I’m saying", 
    "that’s right", "just to be clear", "at this point", "here's the deal", "so anyway"
]

def is_title_case(text):
    return text.istitle()
def is_title(text):
    return text.istitle()
    # Remove leading/trailing spaces
    text = text.strip()

    # Check if the text has a period, which usually means it's not a title
    if "." in text:
        return False
    
    # Split the text into words
    words = text.split()
    
    # Check if the text is too long to be a title
    if len(words) > 10:
        return False
    
    # Count capitalized words
    capitalized_words = sum(1 for word in words if word[0].isupper())
    
    # Consider it a title if at least half the words are capitalized or if the text has special title-like characters
    if capitalized_words >= len(words) / 2 or re.search(r"[:!?]", text):
        return True
    
    return False
def is_title_length(text, min_length=5):
    return len(text) >= min_length
def contains_punctuation(text):
    return any(punct in text for punct in [':', '-', '!', '?'])
def is_text_title(text, min_length=5):
    return (
        text.istitle() and
        len(text) >= min_length and
        contains_punctuation(text)
    )
def is_title_regex(text):
    # Check if text starts with a capital letter and contains a specific pattern
    pattern = r'^[A-Z].*[\:\-].*$'
    return re.match(pattern, text) is not None
def clean_string(text):
    # Step 1: Remove leading numbers and dots
    cleaned_text = re.sub(r'^\d+\.\s*', '', text)

    # Step 2: Replace brackets with a comma (and space)
    cleaned_text = re.sub(r'\((.*?)\)', r', \1', cleaned_text)

    # Step 3: Remove any double spaces or extra spaces around commas
    cleaned_text = re.sub(r'\s*,\s*', ', ', cleaned_text)
    cleaned_text = re.sub(r'\s+', ' ', cleaned_text).strip()

    return cleaned_text
def extract_titles_from_html(html):
    soup = BeautifulSoup(html, 'html.parser')
    titles = []

    # Look for common title tags like <h1>, <h2>, etc.
    for heading in soup.find_all(['h1', 'h2', 'h3']):
        text = heading.get_text()
        if is_title(text):
            titles.append(text)
    
    return titles

def check_text_line_if_title_is(text):
    # 1. Check for any number followed by a period (like "1.", "23.", etc.)
    number_with_period_pattern = r'\d+\.'
    
    # Find all occurrences of numbers followed by a period
    numbers_with_periods = re.findall(number_with_period_pattern, text)
    
    # 2. Check if the string has more uppercase characters than lowercase characters
    num_uppercase = sum(1 for char in text if char.isupper())
    num_lowercase = sum(1 for char in text if char.islower())
    
    more_uppercase_than_lowercase = num_uppercase > num_lowercase

    # 3. Check if the text does not end with a full stop, but ignore cases where it's a number like "1."
    ends_with_fullstop = text.endswith(".")
    
    # Make sure it's not a number followed by a period
    if ends_with_fullstop:
        # Check if the last "sentence" is a number + period
        if re.search(number_with_period_pattern + r'$', text):
            ends_with_fullstop = False  # Ignore if the ending is a number with period

    # 4. Check for the presence of ":" and extract text until that point
    colon_index = text.find(":")
    if colon_index != -1:
        text_until_colon = text[:colon_index + 1]  # Include the colon
    else:
        text_until_colon = None  # No colon found
    
    # 5. Check if all or most words in the text start with uppercase
    words = text.split()
    num_words = len(words)
    num_uppercase_words = sum(1 for word in words if word[0].isupper())
    
    all_words_uppercase = num_uppercase_words == num_words  # All words start with uppercase
    most_words_uppercase = num_uppercase_words > num_words / 2  # More than half start with uppercase

    # Summary of conditions
    collected_info={
         "has_number_with_period": bool(numbers_with_periods),
        "more_uppercase_than_lowercase": more_uppercase_than_lowercase,
        "does_not_end_with_fullstop": not ends_with_fullstop,
        "text_until_colon": text_until_colon,
        "all_words_start_uppercase": all_words_uppercase,
        "most_words_start_uppercase": most_words_uppercase
    }
    is_text_title=False
    if collected_info['more_uppercase_than_lowercase']==True:
        is_text_title=True
        return {
            "is_text_title":is_text_title,
            "text": text
        }
    elif collected_info['has_number_with_period']==True and collected_info['more_uppercase_than_lowercase']==True:
        is_text_title=True
        return {
            "is_text_title":is_text_title,
            "text": text
        }
    elif collected_info['has_number_with_period']==True and collected_info['does_not_end_with_fullstop']==True:
        is_text_title=True
        return {
            "is_text_title":is_text_title,
            "text": text
        }
    elif collected_info['text_until_colon']==True:
        is_text_title=True
        return {
            "is_text_title":is_text_title,
            "text": text_until_colon
        }
    elif collected_info['all_words_start_uppercase']==True or collected_info['most_words_start_uppercase']==True:
        is_text_title=True
        return {
            "is_text_title":is_text_title,
            "text": text_until_colon
        }
    else:
        return {
            "is_text_title":is_text_title,
            "text": text
        }


def extract_keyword_from_sentence(text):
    global nlp
    # Function to extract keywords from a sentence
    def extract_keywords(text):
        doc = nlp(text)
        keywords = [token.text for token in doc if token.pos_ in ["NOUN", "ADJ", "VERB"]]
        return " ".join(keywords)

    # Example sentence to extract keywords from
    sentence = text
    keywords = extract_keywords(sentence)
    return keywords


def is_special_chars_only(s, allowed_chars={'_', '-', ' ', '\n', '\t',''}):
    """
    Check if the string consists only of the specified special characters.
    
    :param s: The string to check.
    :param allowed_chars: A set of allowed characters (default is {'_', '-', ' ', '\n'}).
    :return: True if the string contains only the allowed characters, False otherwise.
    """
    return all(char in allowed_chars for char in s)

def remove_excluded_text(sentence, excluding_list):
    # Create a regex pattern to match excluded words inside brackets
    pattern = r"\[(" + "|".join(map(re.escape, excluding_list)) + r")\]\s*"
    
    # Remove matches from the sentence
    cleaned_sentence = re.sub(pattern, "", sentence, count=1)  # count=1 ensures only the first match is removed
    
    return cleaned_sentence
def plain_text_meta_info(input_file_path):
    excluding_lists = ["[Music].", "[Music]"]
    file_path=input_file_path
    punctuations = ('.', '?', ':', ',', '-')
    # tells only linenumber, if the line title is, and content of the line
    lang='eng'
    try:
        if os.path.exists(file_path):
            with open(file_path, 'r',encoding='utf-8') as file:
                lines = file.readlines()
            my_collection=[]
            if lang=='eng':
                lineExcludes={'_', '-', ' ', '\n', '\t',''}
                for line_number, line in enumerate(lines, start=1):
                    
                    if not is_special_chars_only(line,lineExcludes):
                        keywords_line=extract_keyword_from_sentence(line)
                        check_if_line_is_title=check_text_line_if_title_is(line)
                        is_title=False
                        if check_if_line_is_title['is_text_title']==True:
                            is_title=True
                    
                        cleaned_sentence = remove_excluded_text(line.strip(), excluding_lists)
                        if not cleaned_sentence.strip().endswith(punctuations):
                            cleaned_sentence += ':'
                        line_object = {
                            'linenumber': line_number,
                            'is_title':is_title,
                            'content': cleaned_sentence,
                            'key_words':keywords_line
                        }
                        my_collection.append(line_object)
            return my_collection
        else:
            return "Error: File not found,"+file_path
    except Exception as error:
            return f"Error:There is something wrong! {error}"

def plain_text_only_titles_and_their_contents_and_subconents(input_file_path):
    file_path=input_file_path
   
    # extracts only titles and their content and the content under each title
    lang='eng'
    try:
        if os.path.exists(file_path):
            with open(file_path, 'r',encoding='utf-8') as file:
                lines = file.readlines() 
            my_collection = []
            if lang == 'eng':
                current_title = None
                current_sub_content = []
                
                for line_number, line in enumerate(lines, start=1):
                    line = line.strip()
                    if not is_special_chars_only(line,lineExcludes):
                        check_if_line_is_title=check_text_line_if_title_is(line)
                        keywords_line=extract_keyword_from_sentence(line)
                        if check_if_line_is_title['is_text_title']==True:
                            if current_title is not None:
                                my_collection.append({
                                    'linenumber': current_title['linenumber'],
                                    'is_title': True,
                                    'title_text': current_title['title_text'],
                                    'title_sub_content': ' '.join(current_sub_content),
                                    'key_words':keywords_line
                                })
                            
                            current_title = {
                                'linenumber': line_number,
                                'title_text': line,
                                'key_words':keywords_line
                            }
                            current_sub_content = [] 
                        
                        else:
                            if current_title is not None:
                                current_sub_content.append(line)
                    
                if current_title is not None:
                    my_collection.append({
                        'linenumber': current_title['linenumber'],
                        'is_title': True,
                        'title_text': current_title['title_text'],
                        'title_sub_content': ' '.join(current_sub_content),
                        'key_words':keywords_line
                    })
            else:
                return "Error: Only english language text are supported at this moment."    
            return my_collection
        else:
            return "Error: File not found."+file_path
    except Exception as error:
            return f"Error:There is something wrong! {error}"
def get_titles(input_file_path):
    file_path=input_file_path
    try:
        if os.path.exists(file_path):
            plain_text_titles_info=plain_text_only_titles_and_their_contents_and_subconents(file_path)
            return plain_text_titles_info
        else:
            return "Error: File not found"
    
    except Exception as error:
        return f"Error:There is something wrong! {error}"
def get_titles_as_list(input_file_path):
    file_path=input_file_path
    try:
        if os.path.exists(file_path):
            plain_text_titles_info=plain_text_only_titles_and_their_contents_and_subconents(file_path)
            all_titles_list=[]
            for title in plain_text_titles_info:
                clean_chars = clean_string(title['title_text'])
                clean_list=clean_chars.split(',')
                for chars in clean_list:
                    all_titles_list.append(chars.strip())
            if len(all_titles_list) == 0:
                 structuredPlainText=plain_text_meta_info(input_file_path)
                 if len(structuredPlainText) > 0:
                     for data in structuredPlainText:
                         if data["is_title"]:
                             all_titles_list.append(data["content"])
                         
            if len(all_titles_list) == 0:
                breakdataAtFullstop=break_text_at_fullstop(input_file_path)
                if len(breakdataAtFullstop) > 0:
                    for data in breakdataAtFullstop:       
                        keyword=extract_keyword_from_sentence(data["content"])
                        all_titles_list.append(keyword)
            if len(all_titles_list) > 0:
                exclude_texts=["thank","watching next time","next time","Music"]
                filtered_list = [item.strip() for item in all_titles_list if not any(ex_text in item for ex_text in exclude_texts)]
                all_titles_list = filtered_list
            return all_titles_list
        else:
            return "Error: File not found"
    except Exception as error:
        return f"Error:There is something wrong! {error}"
def get_all_plain_text_info(input_file_path):
        file_path=input_file_path
        try:
            if os.path.exists(file_path):
                plain_text_info=plain_text_meta_info(file_path)
                return plain_text_info
            else:
                return "Error: File not found"
        except Exception as error:
            return f"Error:There is something wrong! {error}"  

def break_text_at_fullstop(input_text_file):
    try:
        # Read text from a file
        with open(input_text_file, 'r') as file:
            text = file.read()

        # Split the text at both full stops ('.') and colons (':')
        sentences = re.split(r'[.:]', text)
        sentences = [sentence.strip() for sentence in sentences if sentence.strip()]
        my_collection=[]
        line_number=0
        for sentence in sentences:
            line_number +=1
            keywords_line=extract_keyword_from_sentence(sentence)
            check_if_line_is_title=check_text_line_if_title_is(sentence)
            is_title=False
            if check_if_line_is_title['is_text_title']==True:
                is_title=True
            line_object = {
                'linenumber': line_number,
                'is_title':is_title,
                'content': sentence.strip(),
                'key_words':keywords_line
            }
            my_collection.append(line_object)
        return my_collection
    except Exception as error:
        return f"Error:There is something wrong! {error}"  
    

def extract_ms_word_text_with_metadata(input_file_path):
    file_path=input_file_path
    try:
        if os.path.exists(file_path):
            document = Document(file_path)
            text_with_formatting = []
            for paragraph in document.paragraphs:
                for run in paragraph.runs:  # Each run is a contiguous block of text with the same formatting
                    text_info = {
                        'text': run.text,
                        'bold': run.bold,
                        'underline': run.underline
                    }
                    text_with_formatting.append(text_info)
            return text_with_formatting
        else:
            return "Error: Word document is not found."
    except Exception as error:
            return f"Error:There is something wrong! {error}"
def extract_ms_word_text(input_file_path):
    file_path=input_file_path
    try:
        if os.path.exists(file_path):
            document = Document(file_path)
            paragraphs_with_formatting = []

            # Iterate through paragraphs
            for line_number, paragraph in enumerate(document.paragraphs, start=1):
                # Initialize a list to hold formatted text segments for this paragraph
                formatted_segments = []

                # Iterate through runs in the paragraph to capture formatting
                for run in paragraph.runs:
                    # Create a dictionary for each run with formatting info
                    run_info = {
                        'text': run.text,
                        'bold': run.bold,
                        'underline': run.underline,
                        'italic': run.italic,
                        'font_size': run.font.size.pt if run.font.size else None,
                        'font_color': run.font.color.rgb if run.font.color else None,
                        'highlight': run.font.highlight_color,
                        'strikethrough': run.font.strike,
                        'superscript': run.font.superscript,
                        'subscript': run.font.subscript,
                        'small_caps': run.font.small_caps,
                        'all_caps': run.font.all_caps,
                    }
                    formatted_segments.append(run_info)

                # Store paragraph information along with line number and formatted segments
                paragraph_info = {
                    'linenumber': line_number,
                    'formatted_segments': formatted_segments
                }
                paragraphs_with_formatting.append(paragraph_info)
            finalized_word_text_info=[]
            for item in paragraphs_with_formatting:
                # print("item line nubmer:",item['formatted_segments'])
                if len(item['formatted_segments']) != 0 :
                    text_collection=[]
                    text_meta_info=[]
                    for segment in item['formatted_segments']:
                        text_collection.append(segment['text']+" ")
                        word_info={
                            'linenumber':item['linenumber'],
                            'word':segment['text'],
                            'Bold':segment['bold'],
                            'Underline':segment['underline'],
                            'Italic':segment['italic'],
                            'font_size':segment['font_size'],
                            'font_color':segment['font_color'],
                            'highlight':segment['highlight'],
                            'strikethrough':segment['strikethrough'],
                            'superscript':segment['superscript'],
                            'subscript':segment['subscript'],
                            'small_caps':segment['small_caps'],
                            'all_caps':segment['all_caps']
                        }
                        text_meta_info.append(word_info)
                   
                    text_line=''.join(text_collection)
                    keywords=extract_keyword_from_sentence(text_line)
                    line_info = {
                        'linenumber': item['linenumber'],
                        'line_text': ''.join(text_collection),
                        'details_meta_info':text_meta_info,
                        'key_words':keywords
                    }
                    finalized_word_text_info.append(line_info)
            return finalized_word_text_info
        else:
            return "Error: Word document did not found."
    except Exception as error:
                return f"Error:There is something wrong! {error}"
def apply_text_shadow_word_document(paragraph, text, font_size,shadow_color):
    # Create a shadow text run (slightly offset and in gray color)
    shadow_run = paragraph.add_run(text)
    shadow_run.font.color.rgb =shadow_color #RGBColor(100, 100, 100)  # Shadow color (grayish)
    shadow_run.font.size = Pt(font_size)
    shadow_run.font.bold = True  # Shadow bold
    # Simulate an offset by adding space between shadow and main text
    paragraph.add_run(" ")
def add_3d_word_document(doc, text):
    # Create a new paragraph
    paragraph = doc.add_paragraph()
    
    # Apply shadow effect first
    apply_text_shadow_word_document(paragraph, text, 48,RGBColor(100, 100, 100))

    # Now add the main text run on top of the shadow
    run = paragraph.add_run(text)
    run.font.bold = True  # Make the text bold to enhance 3D effect
    run.font.size = Pt(48)  # Increase the font size
    run.font.color.rgb = RGBColor(0, 0, 255)  # Set the font color (blue)
def make_3d_word_document(text,save_file):
    # Create a new document
    doc = Document()

    # Add 3D-like text
    add_3d_word_document(doc, text)

    # Save the document
    # doc.save('3d_text.docx')
    doc.save(save_file)
# make_3d_word_document('gide segid','wediseg.docx')

def create_image_with_text(text, image_path, image_width, image_height,
                           image_background_color,
                           x_text_position, y_text_position,text_background_color=None,
                           text_color=None, font_path=None, 
                           font_size=None, shadow_offset=None, shadow_color=None):
    
    # Defaults for optional arguments
    # mytext_background_color=text_background_color if text_background_color else "#FFFFFF"
    mytext_color = text_color if text_color else "#FFFFFF"
    myfont_path = font_path if font_path else "arial.ttf"
    myfont_size = font_size if font_size else 50
    myshadow_offset = shadow_offset if shadow_offset else 0
    myshadow_color = shadow_color if shadow_offset else None
    if text_background_color == None or text_background_color == "None":
        text_background_color = "#FFFFFF"
    print("text_background_color:",text_background_color)
    # Create a new image
    image_size = (int(image_width), int(image_height))
    image = Image.new('RGB', image_size, image_background_color)
    draw = ImageDraw.Draw(image)

    # Load the font
    font = ImageFont.truetype(myfont_path, myfont_size)

    # Calculate the bounding box for the text
    text_bbox = draw.textbbox((x_text_position, y_text_position), text, font=font)

    # Add text background if specified
    if text_background_color:
        draw.rectangle(text_bbox, fill=text_background_color)

    # Add shadow if specified
    if shadow_color:
        draw.text((x_text_position + myshadow_offset, y_text_position + myshadow_offset),
                  text, font=font, fill=myshadow_color)

    # Draw the actual text
    draw.text((x_text_position, y_text_position), text, font=font, fill=mytext_color)

    # Save the image
    image.save(image_path)

def spacy_text(text_file_path, text_to_add_at_title=None,text_to_add_at_content=None,return_result_as_paragraph=None,replacements=None,wraps=None,select_random_text_decorator=None):
    text=plain_text_meta_info(text_file_path)
    # print('text=',text)
    new_text_result=[]
    new_paragraph=""
    for index,item in enumerate(text):
        if item['is_title']==True:
            if text_to_add_at_title !=None:
                if 'where:' in text_to_add_at_title:
                   split_text=text_to_add_at_title.split('where:')
                   #    <break time="500ms"/> where:front
                   first_part=split_text[0]
                   second_part=split_text[1]
                   if second_part=='front':
                        item['content']=first_part+item['content']
                   elif second_part=='end':
                       item['content']=item['content']+first_part
                   elif second_part=='both':
                       item['content']=first_part+item['content']+first_part
                else:
                   item['content']=item['content']+text_to_add_at_title 
        else:
            if text_to_add_at_content !=None:
                if 'where:' in text_to_add_at_content:
                    split_text=text_to_add_at_content.split('where:')
                    first_part=split_text[0]
                    second_part=split_text[1]
                    if second_part=='front':
                        item['content']=first_part+item['content']
                    elif second_part=='end':
                        item['content']=item['content']+first_part
                    elif second_part=='both':
                        item['content']=first_part+item['content']+first_part
                    else:
                        item['content']=item['content']+first_part
                else:
                   item['content']=item['content']+text_to_add_at_content
        if replacements !=None:
            if isinstance(replacements, list):
                for replacements_item in replacements:
                    separation_char=replacements_item.split('=',1)
                    first_part=separation_char[0]
                    second_part=separation_char[1]
                    if 'where:' in second_part:
                        # replacements=[':=<break time="500ms"/> where:front']
                        split_second_part=second_part.split('where:')
                        # where: is not yet implmented
                        if first_part in item['content']:
                            item['content']=item['content'].replace(first_part,split_second_part[0].strip())
                    else:
                        if first_part in item['content']:
                            item['content']=item['content'].replace(first_part,second_part)
            elif isinstance(replacements, str):
                replacing=replacements.split('=')
                first_part=replacing[0]
                second_part=replacing[1]
                item['content']=item['content'].replace(first_part,second_part)
            else:
                print("The variable is neither a list nor a string.")
        if wraps !=None:
            # Find the position of the first '>' and the last '</'
            start_index = wraps.find('>') + 1
            end_index = wraps.rfind('</') 
            if 'where:' in wraps:
                split_wraps=wraps.split('where:')
                first_part=split_wraps[0]
                start_index = first_part.find('>') + 1
                end_index = first_part.rfind('</') 

                second_part=split_wraps[1]
                if str(index) in second_part:
                    # [1,2,3,4]
                    item['content'] = wraps[:start_index] + item['content'] + wraps[end_index:]
                   
            else:
                # Insert the text between the '>' and the dynamic closing tag
                item['content'] = wraps[:start_index] + item['content'] + wraps[end_index:]
        if select_random_text_decorator !=None:
            random_filler = random.choice(filler_words)
            item['content']=random_filler+", "+item['content']
        new_text_result.append(item)
    if return_result_as_paragraph != None:
        if return_result_as_paragraph =='yes':
            for item in new_text_result:
                new_paragraph+=item['content']
            return new_paragraph
        else:
            return new_text_result
    else:
     return new_text_result

# generate question.........................................>
def generate_question_from_sentence_plain_text_using_bert(text_file_path,text_to_add_at_content=None,wraps=None,select_random_text_decorator=None):
    # Load a pre-trained T5 model and tokenizer
    tokenizer = T5Tokenizer.from_pretrained("mrm8488/t5-base-finetuned-question-generation-ap",legacy=True)
    model = T5ForConditionalGeneration.from_pretrained("mrm8488/t5-base-finetuned-question-generation-ap",legacy=True)

    # Input text
    text = plain_text_meta_info(text_file_path)
    new_text_result=[]
    for index,item in enumerate(text):
        # Preprocess text for the model
        # input_text = f"generate question: {text}"
        input_ids = tokenizer(item['content'], return_tensors="pt").input_ids

        # Generate question
        outputs = model.generate(input_ids)
        question = tokenizer.decode(outputs[0], skip_special_tokens=True)
       
        item['question']=question
        if text_to_add_at_content !=None:
                if 'where:' in text_to_add_at_content:
                    split_text=text_to_add_at_content.split('where:')
                    first_part=split_text[0]
                    second_part=split_text[1]
                    if second_part=='front':
                        item['question']=first_part+item['question']
                    elif second_part=='end':
                        item['question']=item['question']+first_part
                    elif second_part=='both':
                        item['question']=first_part+item['question']+first_part
                    else:
                        item['question']=item['question']+first_part
                else:
                   item['question']=item['question']+text_to_add_at_content
        if wraps !=None:
            # Find the position of the first '>' and the last '</'
            start_index = wraps.find('>') + 1
            end_index = wraps.rfind('</') 
            if 'where:' in wraps:
                split_wraps=wraps.split('where:')
                first_part=split_wraps[0]
                start_index = first_part.find('>') + 1
                end_index = first_part.rfind('</') 

                second_part=split_wraps[1]
                if str(index) in second_part:
                    # [1,2,3,4]
                    item['question'] = wraps[:start_index] + item['question'] + wraps[end_index:]
                   
            else:
                # Insert the text between the '>' and the dynamic closing tag
                item['question'] = wraps[:start_index] + item['question'] + wraps[end_index:]
        
        if select_random_text_decorator !=None:
            random_filler = random.choice(filler_words)
            
            item['question']=random_filler+", "+item['question']
        new_text_result.append(item)
    return new_text_result


# with_questions=generate_question_from_sentence_plain_text_using_bert(file_path,select_random_text_decorator='yes')
# print('with_questions:',with_questions[1])

def create_text_file(path,content):
    # Create and open a file for writing
    with open(path, 'w') as file:
        # Write some data to the file
        file.write(content)
    
def create_text_from_docx(filePath, outputFilePath,work_output_directory=None,id=None,name=None):
    try:
        whatTypeOfFile=check_file_type(filePath)
        myoutput="my_input_text_data/"+outputFilePath
        if work_output_directory !=None or work_output_directory !="":
            if name !=None and name !="":
                myoutput=f"{work_output_directory}/{name}_{outputFilePath}" 
            else:
                myoutput=f"{work_output_directory}/{outputFilePath}" 
        else:
            if name !=None and name !="":
                myoutput= myoutput+name+"_"+outputFilePath
            else:
                myoutput= myoutput+outputFilePath
        if whatTypeOfFile["file_type"] == "word_document":
            doc = Document(filePath)
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            # Write the extracted text to a .txt file
            with open(myoutput, 'w', encoding='utf-8') as f:
                f.write("\n".join(text))
            
            print(f"Text successfully written to {myoutput}")
            return myoutput
        else:
            return "Input file is not word document"
    except Exception as error:
        return f"Error: There was error while converting word document to txt.{error}"
def create_txt_from_pdf(filePath,outputFilePath,work_output_directory=None,id=None,name=None):
    try:
            whatTypeOfFile=check_file_type(filePath)
            myoutput="my_input_text_data/"
            if work_output_directory !=None or work_output_directory !="":
                if name !=None and name !="":
                    myoutput=f"{work_output_directory}/{name}_{outputFilePath}" 
                else:
                    myoutput=f"{work_output_directory}/{outputFilePath}"
            else:
                if name !=None and name !="":
                    myoutput= myoutput+name+"_"+outputFilePath
                else:
                   myoutput= myoutput+outputFilePath

            if whatTypeOfFile["file_type"] == "pdf":
                reader = PdfReader(filePath)
                # Extract text and write to a file
                with open(myoutput, "w", encoding="utf-8") as f:
                    for page in reader.pages:
                        f.write(page.extract_text() + "\n")

                return myoutput
            else:
                return f"Input file is not pdf. File is type of {whatTypeOfFile['file_type']}"
    except Exception as error:
        return error
    
def extract_keyword_from_sentence(text):
    # Load spaCy's pre-trained English model
    nlp = spacy.load("en_core_web_sm")

    # Function to extract keywords from a sentence
    def extract_keywords(text):
        doc = nlp(text)
        keywords = [token.text for token in doc if token.pos_ in ["NOUN", "ADJ", "VERB"]]
        return " ".join(keywords)

    # Example sentence to extract keywords from
    # sentence = "A human being standing in a futuristic cityscape with neon lights"
    sentence = text
    keywords = extract_keywords(sentence)
    return keywords

def get_nouns(sentence):

    doc = nlp(sentence)
    noun_phrases = []

    for token in doc:
        if token.pos_ == "NOUN":
            # Collect adjectives modifying the noun
            modifiers = [child.text for child in token.lefts if child.dep_ == "amod"]
            noun_phrase = " ".join(modifiers + [token.text])  # Combine adjectives and noun
            noun_phrases.append(noun_phrase)

    return noun_phrases

def is_tangible_noun(word):
    """
    Check if a noun is a tangible (physical) object using WordNet.
    """
    # print("word:",word)
    word = lemmatizer.lemmatize(word)  # Convert to singular form
    synsets = wordnet.synsets(word, pos=wordnet.NOUN)  # Get all noun synsets

    # print(f"Checking word: {word}, Synsets: {[syn.name() for syn in synsets]}")  # Debugging output
    
    if not synsets:
        return False  # If no synsets exist, the word is likely not in WordNet

    # Categories considered tangible
    tangible_categories = {'physical_entity', 'object', 'food', 'artifact', 'place', 'substance'}

    for syn in synsets:
        hypernyms = [lemma.name().split(".")[0] for lemma in syn.closure(lambda s: s.hypernyms())]  # Get all hypernyms
       
        # Extract the lemma (base word) from synset (Fixes the ".n.01" issue)
        lemma_names = {lemma.name() for synset in synsets for lemma in synset.lemmas()}
        
        # print("lemma_names:",lemma_names)
        # If any hypernym matches tangible categories, the word is tangible
        if any(h in tangible_categories for h in hypernyms):
            return True

    return False

def filter_tangible_nouns(word_list):
    """
    Filters only tangible (touchable/seeable) nouns and removes duplicates.
    """
    processed_nouns = set()  # Use a set to remove duplicates
    doc = nlp(" ".join(word_list))  # Process words with spaCy
    for chunk in doc.noun_chunks:  # Get noun phrases
        phrase = chunk.text.lower()  # Convert to lowercase for consistency
        words = phrase.split()  # Split phrase into individual words
        # If at least one word in the phrase is tangible, keep the whole phrase
        if any(is_tangible_noun(word) for word in words):
            processed_nouns.add(phrase)  # Add noun phrase to set

    modifiedProcessedNouns=[]
    if len(list(processed_nouns)) > 0:
        for noun in list(processed_nouns):
            mynoun=noun.split()
            if len(mynoun) > 3:
                pass
            else:
                modifiedProcessedNouns.append(noun)

    return modifiedProcessedNouns  # Convert back to list
def remove_partial_matches(list1, list2):
    return [item for item in list1 if not any(sub in item for sub in list2)]

def getItemsForDownload(filePath):
    # "my_input_text_data/LMLSISU0GFc_transcribed.txt"
    text=plain_text_meta_info(filePath)
    itemsForDownload=[]
    for line in text:
        nouns=get_nouns(line['content'])
        if len(nouns) > 0:
            for noun in nouns:
                itemsForDownload.append(noun)
    modifiedNouns=[]
    if len(itemsForDownload)>0:
        filtered_nouns = filter_tangible_nouns(itemsForDownload)
        modifiedNouns=filtered_nouns

    # if len(modifiedNouns) > 0:
    #     list2=['a.','i.','you.','you','i','a','the.','the','top','bottom','right','left','side',"video","videos","music","[Music].","morning","after","before","back","comment","share","subscribe","channel","my","your","his","her","hello","guy","guys","soft","simple","hard","big","small","little","bit","very"]
    #     filtered_list = remove_partial_matches(modifiedNouns, list2)
    #     print("filtered_list::::????",filtered_list)
    #     modifiedNouns=filtered_list
    # print("modifiedNouns::::",modifiedNouns)
    return modifiedNouns


def modifyInputTextForProcess(filePath,processIsFor):
    try:
        print("filePath:::",filePath)
        file_content=plain_text_meta_info(filePath)
        modifiedTextContent=[]
        for line in file_content:
            print("line:",line)
            linenumber=line['linenumber']
            is_title=line['is_title']
            if is_title:
                cleaned_text = re.sub(r'[^a-zA-Z0-9\s]', '', line['content'])
                splitContent=cleaned_text.split()
                # Join all words except the last one with a comma
                result = ", ".join(splitContent[:-1])

                # Add the last word followed by three commas
                result += f", {splitContent[-1]},,,,"
                content=result
            else:
                splitContent=line['content'].split()
                # Join the words with a comma, except for the last word
                result = ",, ".join(splitContent[:-1]) + " " + splitContent[-1]
                content=result
            
            modifyContent=checkForNumberInSentence(content)
            line['content']=modifyContent
            modifiedTextContent.append(line)
        modifyTextFile=saveModificationToFile(filePath,processIsFor,modifiedTextContent)
        if 'Error' in modifyTextFile:
            print(f"Error: There was error modifying input text file.{modifyTextFile}")
            return f"Error: There was error modifying text input as ai text audio"
        else:
            print("Input text file modified.")
            return modifyTextFile
    except Exception as error:
        return f"Error: There was error modifying text input as ai text audio"


def checkForNumberInSentence(sentence):
    text=sentence.split()
    modifiedText=[]
    for item in text:
        modifiedItem=item
        isNumber=is_number(item)
        if isNumber:
            modifiedItem=number_to_words(item) 
        else:
            isItemCurrenctWithNumber=is_currency_number(item)
            isCurrencyWithNumber=currency_with_number_to_words(item)
            if isCurrencyWithNumber:
                modifiedItem=isCurrencyWithNumber
            else:
                sparateTextAndNumber=separate_text_number(item) 
                if len(sparateTextAndNumber) == 1:
                    modifiedItem=number_to_words(sparateTextAndNumber[0])
                else:
                    firstSplit=number_to_words(sparateTextAndNumber[0])
                    secondSplit=number_to_words(sparateTextAndNumber[1])
                    modifiedItem=f"{firstSplit}{secondSplit}"
       
        modifiedText.append(modifiedItem)
    return ' '.join(modifiedText)

def saveModificationToFile(filePath,processIsFor,data):
    try:
        checkFilePath=check_file_type(filePath)
        fileName=f"text_input_for_ai_on_{processIsFor}.txt"
        fullPath=combine_path(checkFilePath["directory_path"],fileName) 
        # Find the highest line number to create enough lines
        max_line = max(item['linenumber'] for item in data)

        # Initialize a list with empty lines
        lines = [""] * max_line

        # Populate the list with content at the correct line number
        for item in data:
            line_index = item['linenumber'] - 1  # Convert to 0-based index
            lines[line_index] = item['content']
        # Save to a text file
        with open(fullPath, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        return fullPath
    except Exception as error:
        return f"Error:There was error while saving modified text.{error}"

def getTags(filePath):
    try:
        with open(filePath, "r", encoding="utf-8") as file:
            text = file.read()
        doc = nlp(text)
        entities = [(ent.text, ent.label_) for ent in doc.ents]
        flat_list = [item for sublist in entities for item in sublist]
        return flat_list
    except Exception as error:
        return f"Error:There was error while extracting tags from text. {error}"

def extract_description(inputText):
    try:
        text=""
        if os.path.exists(inputText):
            with open(inputText, "r", encoding="utf-8") as file:
                text = file.read()
        else:
            text=inputText
        if text !="":
            # Try to extract description from a structured format
            match = re.search(r"(?i)Description:\s*(.+)", text, re.DOTALL)
            if match:
                description = match.group(1).strip()
            else:
                # If no structured description, take the first few sentences
                sentences = re.split(r'(?<=[.!?]) +', text)  # Split into sentences
                description = " ".join(sentences[:3])  # Take the first 3 sentences
            text_cleaned=clean_text_on_extracted_description(description.strip())
            return text_cleaned
        else:
            return text
    
    except Exception as error:
        return f"Error: Could not extract description. {error}"

def clean_text_on_extracted_description(text):
    # Remove markdown headers (##, ###, etc.)
    text = re.sub(r'#+\s*', '', text)
    
    # Remove horizontal rules like '---' or '___'
    text = re.sub(r'[-_]{3,}', '', text)
    
    # Replace multiple newlines with a single newline
    text = re.sub(r'\n\s*\n+', '\n', text)
    
    # Remove any extra whitespace from start/end of lines
    text = '\n'.join(line.strip() for line in text.strip().splitlines())
    
    # Optionally remove emojis (if needed)
    # text = re.sub(r'[^\w\s.,!?\'\"-]', '', text)

    return text
def extract_summary(inputText, num_sentences=300):
    try:
        text=""
        if os.path.exists(inputText):
            with open(inputText, "r", encoding="utf-8") as file:
                text = file.read()
        else:
            text=inputText
        if text !="":
            # Remove unnecessary newlines and extra spaces
            text = re.sub(r'\s+', ' ', text).strip()

            # Summarization using sumy (TextRank algorithm)
            parser = PlaintextParser.from_string(text, Tokenizer("english"))
            summarizer = TextRankSummarizer()
            summary = summarizer(parser.document, num_sentences)

            # Convert summarized sentences into a single description string
            description = " ".join(str(sentence) for sentence in summary)
            return description.strip()[:4000]
        else:
            return text
    except Exception as error:
        return f"Error: Could not summarize description. {error}"


def extract_title(inputText):
    try:
        text=""
        if os.path.exists(inputText):
            with open(inputText, "r", encoding="utf-8") as file:
                text = file.read()
        else:
            text=inputText
        
        if text !="":
            # Remove unnecessary newlines and extra spaces
            text = re.sub(r'\s+', ' ', text).strip()

            # Summarization using sumy (TextRank algorithm)
            parser = PlaintextParser.from_string(text, Tokenizer("english"))
            summarizer = TextRankSummarizer()
            summary = summarizer(parser.document, 1)  # Extract only 1 key sentence

            # Convert the extracted sentence into a title
            title = str(summary[0]) if summary else ""
            print("title:::",title)
            return title.strip()[:90]  # Limit to 90 characters (YouTube title limit)
        else:
            return text

    except Exception as error:
        return f"Error: Could not generate title. {error}"

def extract_tags(inputText, num_sentences=3, max_tags=10):
    try:
        text=""
        if os.path.exists(inputText):
            with open(inputText, "r", encoding="utf-8") as file:
                text = file.read()
        else:
            text=inputText

        if text !="":
            # Load spaCy NLP model
            nlp = spacy.load("en_core_web_sm")
            # Remove unnecessary spaces
            text = re.sub(r'\s+', ' ', text).strip()

            # Summarization using sumy (TextRank algorithm)
            parser = PlaintextParser.from_string(text, Tokenizer("english"))
            summarizer = TextRankSummarizer()
            summary = summarizer(parser.document, num_sentences)  # Summarize

            # Convert summarized sentences into a string
            summary_text = " ".join(str(sentence) for sentence in summary)

            # Extract keywords (nouns, proper nouns) using spaCy
            doc = nlp(summary_text)
            keywords = [token.text for token in doc if token.pos_ in ["NOUN", "PROPN"]]

            # Remove duplicates and limit to max_tags
            tags = list(set(keywords))[:max_tags]

            # Optional: Pre-filter text using POS tagging (nouns & adjectives only)
            filtered_text = " ".join([token.text for token in doc if token.pos_ in ("NOUN", "PROPN", "ADJ")])

            # Now feed filtered text into KeyBERT
            kw_model = KeyBERT()
            tags = kw_model.extract_keywords(filtered_text, keyphrase_ngram_range=(1, 3), stop_words='english', top_n=10)
            mytags=[]
            # Display tags
            for tag in tags:
                # print("tag[0]::",tag[0])
                mytags.append(tag[0])

            return mytags
        else:
            return text

    except Exception as error:
        return f"Error: Could not extract tags. {error}"


def translation(text_to_translate,fromLanguage,toLanguage):
    
    try:
        translated_text = GoogleTranslator(source=fromLanguage, target=toLanguage).translate(text_to_translate)
        # englishAbbrivationInTigrina
        if fromLanguage !="ti":
            checkForEnglishWordsInText = replace_english_words_with_tigrigna(translated_text)
            obj={
                "text":text_to_translate,
                "translated":checkForEnglishWordsInText
            }
            return obj
        else:
            obj={
                "text":text_to_translate,
                "translated":translated_text
            }
            return obj
        
    except Exception as error:
          return f"Error: {error}"

def translateText(fromLanguage,toLanguage,filePath=None,text=None):
    try:
        collection=[]
        if filePath !=None:
            getInfo=get_all_plain_text_info(filePath)
            for line in getInfo:
                content=line["content"]
                trans=translation(content,fromLanguage,toLanguage)
                if "Error" in trans:
                    print(trans)
                else:
                    collection.append(trans)
        
        if text != None:
            trans=translation(text,fromLanguage,toLanguage)
            collection.append(trans)
        return collection
    except Exception as error:
        return f"Error: There was error on translation.{error}"


def generateTitleFromBigText(inputText):
    try:
        text=""
        summarizer=""
        if os.path.exists(inputText):
            # Load text
            with open(inputText, "r", encoding="utf-8") as f:
                text = f.read()
        else:
            text=inputText
        if text !="":
           # Use a small T5 model for summarization
            summarizer = pipeline("summarization", model="t5-small")

            # Token count or character-based heuristic
            input_text = text[:1000]  # still smart to truncate long inputs

            # Avoid warning by adapting max_length
            word_count = len(input_text.split())
            max_len = min(20, word_count * 2)  # adjust if input is short
            min_len = min(5, word_count)       # avoid setting min > word count

            # Generate summary
            summary = summarizer(
                input_text,
                max_length=max_len,
                min_length=min_len,
                do_sample=False
            )

            title = summary[0]['summary_text']
            return title
        else:
            return text
        

    except Exception as error:
        return f"Error: {error}"

current_text="C:\\Users\\GideSegid\\Desktop\\test\\current_project\\current.txt"



def extract_title_with_keywords(inputText):
    try:
        # Step 1: Read the input
        if os.path.exists(inputText):
            with open(inputText, "r", encoding="utf-8") as file:
                text = file.read()
        else:
            text = inputText

        # Step 2: Clean the text
        text = re.sub(r'\s+', ' ', text).strip()

        
        keywords = kw_model.extract_keywords(text, keyphrase_ngram_range=(1, 2), stop_words='english', top_n=5)
        top_keywords = [kw[0] for kw in keywords]
        # print("Top keywords:", top_keywords)

        # Step 4: Generate a basic title
        title = " ".join(top_keywords).title()
       
        # Step 5: Optional paraphrasing to make it more natural
        # Uncomment the below if you want paraphrasing (requires model download)
        paraphraser = pipeline("text2text-generation", model="Vamsi/T5_Paraphrase_Paws")
        paraphrased = paraphraser(f"paraphrase: {title}", max_length=60, num_return_sequences=1)
        title = paraphrased[0]['generated_text']
        # print("title/////:",title)
        return top_keywords  # Limit title length

    except Exception as e:
        return f"Error: Could not generate title. {e}"
    



def format_title_with_templates(inputText,generalIdeaOfText=None):
    try:
        keywords = extract_title_with_keywords(inputText)
        general_templates = [
            "How {0} Are Changing the Game in {1}",
            "Top 5 {0} You Should Know About in {1}",
            "What You Need to Know About {0} and {1}",
            "The Future of {0}: What {1} Tells Us",
            "Are {0} the Key to {1}?",
            "Why {0} Could Be the Future of {1}"
        ]
        motivation_templates = [
            "How {0} Can Transform Your {1}",
            "Top 5 {0} to Supercharge Your {1}",
            "What You Need to Know About {0} and Finding {1}",
            "The Future of {0}: What It Means for Your {1}",
            "Are {0} the Secret to Unlocking {1}?",
            "Why {0} Could Be the Catalyst for Your {1}"
        ]
        technology_templates = [
            "How {0} Is Disrupting the {1} Industry",
            "Top 5 Innovations in {0} Transforming {1}",
            "What {0} Means for the Future of {1}",
            "The Rise of {0}: A New Era for {1}",
            "Is {0} the Next Big Thing in {1}?",
            "Why {0} Matters More Than Ever in {1}",
            "How {0} Is Redefining {1} as We Know It",
            "What Experts Are Saying About {0} in {1}",
            "The Role of {0} in Shaping the Future of {1}",
            "Could {0} Be the Solution to {1}'s Biggest Challenges?"
        ]
        software_templates = [
            "How {0} Software Is Revolutionizing {1}",
            "Top 5 {0} Tools You Need for Better {1}",
            "The Best {0} Software to Improve Your {1}",
            "Why {0} Software Is a Game-Changer for {1}",
            "Is {0} the Ultimate Solution for {1} Challenges?",
            "Everything You Should Know About {0} Software in {1}",
            "How {0} Platforms Are Shaping the Future of {1}",
            "What Makes {0} Software Stand Out in {1}",
            "The Evolution of {0}: What It Means for {1}",
            "Could {0} Be the Software That Transforms {1}?"
        ]
        computer_templates = [
            "How {0} Computers Are Changing the Way We {1}",
            "The Evolution of {0} Computers in the World of {1}",
            "Top 5 {0} Computers That Excel in {1}",
            "Why {0} Computing Power Matters for {1}",
            "The Role of {0} Computers in Modern {1}",
            "Is {0} the Future of Personal Computing in {1}?"
        ]
        hardware_templates = [
            "How {0} Hardware Is Transforming {1} Performance",
            "Top 5 {0} Hardware Upgrades to Boost Your {1}",
            "Why {0} Hardware Is Critical for {1} Applications",
            "The Impact of {0} Components on {1}",
            "What You Should Know About {0} Hardware and {1}",
            "Could {0} Hardware Be the Missing Piece in {1}?"
        ]
        network_templates = [
            "How {0} Networks Are Powering the Future of {1}",
            "Top 5 {0} Networking Technologies Shaping {1}",
            "Why {0} Network Infrastructure Is Essential for {1}",
            "The Role of {0} Networks in Modern-Day {1}",
            "What Makes {0} Networking Crucial to {1}",
            "Is {0} the Key to More Secure {1} Networks?"
        ]
        internet_templates = [
            "How {0} Is Changing the Way We Use the Internet for {1}",
            "Top 5 Internet Innovations in {0} for Better {1}",
            "Why {0} Is Reshaping the Internet and {1}",
            "The Future of Internet {0}: What It Means for {1}",
            "What {0} Tells Us About the Internet’s Role in {1}",
            "Is {0} the Future of the Internet in {1}?"
        ]
        bitcoin_templates = [
            "How Bitcoin Is Disrupting the {0} Industry",
            "Top 5 Things to Know About Bitcoin and {0}",
            "Why Bitcoin Could Be the Future of {0}",
            "What Bitcoin’s Rise Means for {0}",
            "Is Bitcoin the Key to a New Era of {0}?",
            "How Bitcoin Is Changing the Way We Think About {0}",
            "The Impact of Bitcoin on {0}: What You Need to Know",
            "Bitcoin vs. Traditional {0}: What’s the Difference?",
            "What Experts Are Saying About Bitcoin in {0}",
            "Could Bitcoin Solve the Biggest Problems in {0}?"
        ]
        growth_template=[# 🔥 Growth-Oriented
            "How {0} Can Help You Grow in {1}",
            "The Habits That Took My {1} to the Next Level with {0}",
            "Why {0} Is the Secret Weapon for Growth in {1}",
            "The #1 Growth Hack Using {0} in Your {1}",
            "What {0} Taught Me About Leveling Up in {1}",
            "From Stuck to Thriving: Using {0} in {1}",
            "How Successful People Use {0} to Dominate {1}"
            ]
        education_template=[
            "Everything You Need to Know About {0} in {1}",
            "The Most Underrated Learning Method: {0} in {1}",
            "Learn {0} in Just 10 Minutes — Master {1}",
            "Why Learning {0} Is a Game-Changer in {1}",
            "What School Didn’t Teach You About {0} and {1}",
            "The Science Behind {0} in {1} Education",
            "How {0} Makes Learning {1} 10x Easier",]
        health_template=[
            "The Surprising Health Benefits of {0} in Your {1}",
            "Why Your {1} Needs More of {0} (Backed by Science)",
            "This One Change with {0} Transformed My {1}",
            "Are You Doing {0} Right? Boost Your {1} Naturally",
            "{0} for a Healthier, Stronger {1}",
            "How {0} Can Improve Mental Health in {1}",
            "The Link Between {0} and Lasting Wellness in {1}",
        ]
        personal_development_template=[
            # ✨ Personal Development Blend
            "Why {0} Might Be the Key to Unlocking Your Best {1}",
            "How {0} Builds Discipline in {1}",
            "{0} Changed My Perspective on {1}",
            "What Happens When You Stay Consistent with {0} in {1}",
            "The Mindset Shift Behind {0} That Boosts Your {1}",
            "How to Rewire Your Brain with {0} in {1}"
        ]
        templates=""
        
        if generalIdeaOfText !=None:
            if generalIdeaOfText =="general":
                templates=general_templates
            elif generalIdeaOfText=="growth":
                templates=growth_template
            elif generalIdeaOfText=="education":
                templates=education_template
            elif generalIdeaOfText=="health":
                templates=health_template
            elif generalIdeaOfText=="personalDevelopment":
                templates=personal_development_template
            elif generalIdeaOfText=="technology":
                templates=technology_templates
            elif generalIdeaOfText=="software":
                templates=software_templates
            elif generalIdeaOfText=="computer":
                templates=computer_templates
            elif generalIdeaOfText=="hardware":
                templates=hardware_templates
            elif generalIdeaOfText=="network":
                templates=network_templates
            elif generalIdeaOfText=="internet":
                templates=internet_templates
            elif generalIdeaOfText=="bitcoin":
                templates=bitcoin_templates
            elif generalIdeaOfText=="motivation":
                template=motivation_templates
            else:
                templates=general_templates
        else:
            templates=general_templates
            
        # Basic fallback keywords if none are found
        fallback_keywords = ["Innovation", "Education", "Growth", "Health", "Learning"]
        
        # Ensure at least 2 keywords
        if not keywords or len(keywords) < 2:
            keywords = (keywords if keywords else []) + random.sample(fallback_keywords, 2 - len(keywords))

        random.shuffle(templates)
        titles = []

        for _ in range(3):
            # If keywords < 2 after all fallback attempts, duplicate the only keyword
            kw_sample = random.sample(keywords, 2) if len(keywords) >= 2 else [keywords[0], keywords[0]]
            template = random.choice(templates)
            raw_title = template.format(*kw_sample)

            # Grammar clean-up using spaCy
            doc = nlp(raw_title)
            cleaned = doc.text  # or just use raw_title if no processing is needed
            titles.append(cleaned)

        return titles
    except Exception as error:
        return f"Error: There was error while making a topic or title.{error}"

def get_synset(word):
    synsets = wordnet.synsets(word)
    if synsets:
        return synsets[0].definition()
    return None

def lemmatize_and_expand(keywords):
    return {word: get_synset(lemmatizer.lemmatize(word)) for word in keywords}

def extract_keywords(text, num_keywords=10):
    keywords = kw_model.extract_keywords(text, keyphrase_ngram_range=(1, 2), stop_words='english', top_n=num_keywords)
    return [kw[0] for kw in keywords]

def transformer_summary(text):
    # Handle large input by chunking
    return summarizer_bart_facebook(text[:1024], max_length=130, min_length=30, do_sample=False)[0]['summary_text']


def get_core_idea(inputText):
    try:
        text = ""
        if os.path.exists(inputText):
            with open(inputText, "r", encoding="utf-8") as f:
                text = f.read()
        else:
            text=inputText
        summary = transformer_summary(text)  # or summarize_text(text)
        keywords = extract_keywords(summary)
        concepts = lemmatize_and_expand(keywords)

        return {
            "summary": summary,
            "keywords": keywords,
            "concepts": concepts
        }
    except Exception as error:
        return f"Error: There was error while creating core idea from text.{error}"

def format_core_idea_to_create_full_description(data):
    output = ""

    # Summary
    output += f"{data['summary']}\n\n"

    # Keywords
    output += "Keywords:\n"
    for idx, kw in enumerate(data['keywords'], start=1):
        output += f"{idx}. {kw}\n"

    # Concepts
    output += "\nConcepts:\n"
    for field, value in data['concepts'].items():
        output += f"{field}: {value}\n"
    output += "\nTags:\n"
    for tag in data['tags']:
        output += f"#{tag} "
    return output

def correct_grammar(text):
    model_name = "vennify/t5-base-grammar-correction"
    tokenizer = T5Tokenizer.from_pretrained(model_name,legacy=True)
    model = T5ForConditionalGeneration.from_pretrained(model_name)
    input_text = f"grammar: {text}"
    input_ids = tokenizer.encode(input_text, return_tensors="pt")
    outputs = model.generate(input_ids, max_length=64)
    corrected = tokenizer.decode(outputs[0], skip_special_tokens=True)
    return corrected

def generate_title_transformer(summary):
    generator = pipeline("text2text-generation", model="t5-small")
    prompt = f"generate a catchy title: {summary}"
    result = generator(prompt, max_length=20, num_return_sequences=1)
    return result[0]['generated_text']

def encode_text_to_one_line(text):
     return text.replace('\n', '\\n')

def decode_text_to_one_line(encoded_text):
    return encoded_text.replace('\\n', '\n')

try:
    from transformers import pipeline
    use_transformers = True
except ImportError:
    use_transformers = False

# Sample emojis per theme
EMOJIS = {
    "clickbait": ["😱", "🔥", "🚀", "🤯", "👀"],
    "professional": ["📘", "📊", "💡", "🧠", "🔍"],
    "casual": ["✨", "😉", "🤓", "🙌", "🎉"]
}

def extract_keywords(summary):
    words = re.findall(r'\b\w+\b', summary.lower())
    ignore = {"the", "is", "and", "of", "to", "in", "a", "this", "for", "on", "by", "with"}
    keywords = [w for w in words if w not in ignore and len(w) > 3]
    keywords.sort(key=len, reverse=True)
    return keywords[:3]

def generate_title_hook(summary, style="clickbait", use_ai=False, max_length=100):
    if summary !="":
        keywords = extract_keywords(summary)
        print("keywords:",keywords)
        main_kw = keywords[0].capitalize() if keywords else "This Topic"

        # Hooks
        hooks = {
            "clickbait": [
                f"You Won’t Believe What Happens With {main_kw}",
                f"This Changes Everything: {main_kw}",
                f"What No One Told You About {main_kw}",
                f"Why {main_kw} Is Blowing Up 🔥"
            ],
            "professional": [
                f"An In-Depth Guide to {main_kw}",
                f"Understanding the Power of {main_kw}",
                f"Everything You Should Know About {main_kw}",
                f"{main_kw}: A Complete Overview"
            ],
            "casual": [
                f"Let’s Talk About {main_kw}!",
                f"Here’s the Scoop on {main_kw}",
                f"Why We’re Loving {main_kw} Right Now",
                f"{main_kw} Just Got More Interesting"
            ]
        }

        if use_ai and use_transformers:
            gen = pipeline("text2text-generation", model="t5-small")
            prompt = f"generate a catchy title: {summary}"
            result = gen(prompt, max_length=20)[0]["generated_text"]
            return result.strip()[:max_length]

        # Choose title from hooks
        title = random.choice(hooks.get(style, hooks["casual"]))
        
        # Add emoji if appropriate
        emoji = random.choice(EMOJIS.get(style, []))
        if style in ["clickbait", "casual"]:
            title = f"{emoji} {title}"

        return title[:max_length]
    else:
        return ""


# clickbaitTitle=generate_title(title, style="clickbait")
# print("clickbaitTitle:",clickbaitTitle)
# professionalTitle=generate_title(title, style="professional")
# print("professionalTitle:",professionalTitle)
# casualTitle=generate_title(title, style="casual")
# print("casualTitle:",casualTitle)
# aiBasedTitle=generate_title(title, use_ai=True)
# print("aiBasedTitle:",aiBasedTitle)

def categorizeTextPerTime(inputText,chunk_seconds):
    # Settings
    words_per_minute = 150
    words_per_second = words_per_minute / 60
    # chunk_start_times = [0, 5, 15]  # in seconds
    text=""
    if os.path.exists(inputText):
        # Read the text file
        with open(inputText,  'r', encoding='utf-8') as file:
            text = file.read()

    # Preprocess: remove newlines and extra spaces
    text = re.sub(r'\s+', ' ', text).strip()

    # Split into words
    words = text.split()
    word_count = len(words)

    chunks = []
    word_index = 0

    for i in range(len(chunk_seconds)):
        start_time = chunk_seconds[i]
        end_time = chunk_seconds[i+1] if i+1 < len(chunk_seconds) else None

        # Estimate how many words to include
        if end_time is not None:
            duration = end_time - start_time
            est_word_count = int(duration * words_per_second)
        else:
            est_word_count = len(words) - word_index  # remainder

        # Start collecting words
        segment_words = words[word_index:word_index + est_word_count]
        sentence_end_found = False
        j = word_index + est_word_count

        # Extend until we find a sentence end
        while j < len(words):
            segment_words.append(words[j])
            if re.search(r'[.!?]$', words[j]):
                j += 1
                sentence_end_found = True
                break
            j += 1

        word_index = j  # move to next starting point
        text=' '.join(segment_words)
        summary=transformer_summary(text)
        translateText=translation(summary,"en","ti")
        chunks.append({
            "start": f"{start_time:.2f}",
            "end": f"{end_time:.2f}" if end_time is not None else "END",
            "text": ' '.join(segment_words),
            "summary":summary,
            # "tigrina":translateText["translated"]
        })
  
    summaryInfo=[]
    for items in chunks:
        start=int(items["start"].split('.')[0]) / 60
        end=items["end"]
        if items["end"] !="END":
            end=int(items["end"].split('.')[0]) / 60

        summaryText=f"{start}-{end}:{items['summary']}"
        summaryInfo.append(summaryText)
    
    # Output JSON
    json_output = json.dumps(chunks, indent=4)
    description = "Chapters:\n" + "\n".join(f"- {item}" for item in summaryInfo)
    return description

# inputText="C:\\Users\\GideSegid\\Desktop\\test\\current_project\\current.txt"
# chunk_seconds = [30, 60, 90,120,150,180,210,240]
# categorize=categorizeTextPerTime(inputText,chunk_seconds)
# print("categorize:",categorize)


def youtubeTagsCleanUp(raw_tags):
    try:
        # Step 1: Split and clean each tag
        tags = [tag.strip() for tag in raw_tags.split(',') if tag.strip()]

        # Step 2: Remove duplicates while preserving order
        seen = set()
        unique_tags = []
        for tag in tags:
            if tag not in seen:
                seen.add(tag)
                unique_tags.append(tag)

        # Step 3: Limit to 500 characters total including separators
        final_tags = []
        total_length = 0
        for i, tag in enumerate(unique_tags):
            tag_length = len(tag)
            separator_length = 2 if i > 0 else 0  # ", " between tags

            if total_length + tag_length + separator_length <= 500:
                final_tags.append(tag)
                total_length += tag_length + separator_length
            else:
                break

        return final_tags
    except Exception as error:
        return f"Error: There was an error while cleaning up the YouTube tags. {error}"


def youtubeDescriptionCleanUp(text):
    try:
        # Normalize unicode (e.g., accents, weird encodings)
        text = unicodedata.normalize('NFKC', text)
        
        # Remove control characters and non-printable characters
        text = ''.join(c for c in text if c.isprintable())

        # Optional: Remove emojis (or keep them if you want)
        emoji_pattern = re.compile("["
            u"\U0001F600-\U0001F64F"  # emoticons
            u"\U0001F300-\U0001F5FF"  # symbols & pictographs
            u"\U0001F680-\U0001F6FF"  # transport & map symbols
            u"\U0001F1E0-\U0001F1FF"  # flags (iOS)
            "]+", flags=re.UNICODE)
        text = emoji_pattern.sub('', text)

        # Collapse multiple spaces/newlines into clean structure
        text = re.sub(r'[ \t]+', ' ', text)           # multiple spaces → 1 space
        text = re.sub(r'\n{3,}', '\n\n', text)        # max 2 line breaks in a row
        text = text.strip()

        # Optional: Remove suspicious HTML or script-like text
        text = re.sub(r'<[^>]+>', '', text)           # strip HTML tags
        text = re.sub(r'(javascript:|onerror=|onload=)', '', text, flags=re.IGNORECASE)

        # Limit to 5000 characters (YouTube description limit)
        return text[:5000]
    except Exception as error:
        return f"Error: There was error while cleaning up youtube descriptions.{error}"